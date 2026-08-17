import os
import json
import uuid
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity, get_jwt
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, timedelta, timezone, date
import re
import base64
import functools
import email_service
import google_sheets_service
import microsoft_excel_service
import google_forms_api
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__, static_folder='../dist', static_url_path='/')
CORS(app)

# Configurations
app.config['JWT_SECRET_KEY'] = os.environ.get('JWT_SECRET_KEY', 'default-secret-key-fallback')
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(days=7)
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024 # 20MB max upload
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(os.path.dirname(__file__), 'instance', 'database.sqlite')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

jwt = JWTManager(app)
from extensions import db
db.init_app(app)

from models import User, Intern, BlacklistedIntern, Attendance, DocumentRequest, Form, FormSubmission, DocumentTemplate, DocumentLifecycle, Notification, SystemLog, Message, AuditEvent


# Improve SQLite concurrency and performance using WAL mode
from sqlalchemy import event
from sqlalchemy.engine import Engine

@event.listens_for(Engine, "connect")
def set_sqlite_pragma(dbapi_connection, connection_record):
    # Only apply if it's an sqlite3 connection
    if type(dbapi_connection).__module__ == "sqlite3":
        cursor = dbapi_connection.cursor()
        try:
            cursor.execute("PRAGMA journal_mode=WAL")
            cursor.execute("PRAGMA synchronous=NORMAL")
        except Exception as e:
            # Fallback for WSL2/Docker Windows volume mounts where WAL fails with disk I/O error
            print(f"WSL2/Docker Volume Mount detected (WAL mode failed: {e}). Falling back to journal_mode=DELETE")
            try:
                cursor.execute("PRAGMA journal_mode=DELETE")
                cursor.execute("PRAGMA synchronous=FULL")
            except Exception:
                pass
        
        try:
            cursor.execute("PRAGMA cache_size=-64000")  # 64MB cache
            cursor.execute("PRAGMA temp_store=MEMORY")
            cursor.execute("PRAGMA mmap_size=30000000000")
        except Exception:
            pass
        cursor.close()

# Allow download endpoints (opened via window.open) to authenticate using ?token=...
# in addition to the standard Authorization: Bearer header.
app.config['JWT_TOKEN_LOCATION'] = ['headers', 'query_string']
app.config['JWT_QUERY_STRING_NAME'] = 'token'

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def _ext_for_file_type(file_type: str) -> list:
    """Return the list of allowed extensions for the given logical file_type."""
    return FILE_TYPE_EXTENSIONS.get((file_type or '').strip().lower(), FILE_TYPE_EXTENSIONS['any'])


def _allowed_file(file_name: str, file_type: str) -> bool:
    ft = (file_type or '').strip().lower() or 'any'
    return any(file_name.lower().endswith(ext) for ext in _ext_for_file_type(ft))


def _file_extension(file_name: str, file_type: str) -> str:
    """Pick the canonical extension (incl. leading '.') for an upload."""
    ft = (file_type or '').strip().lower() or 'any'
    low = (file_name or '').lower()
    for ext in _ext_for_file_type(ft):
        if low.endswith(ext):
            return ext
    if '.' in low:
        cand = '.' + low.rsplit('.', 1)[1]
        return cand if 1 < len(cand) <= 10 else '.pdf'
    return '.pdf'


def friendly_doc_filename(intern_id: int, doc_id, label: str, file_type: str) -> str:
    """Build an IT-friendly filename: intern-{id}-{doc_id}-{safe-label}.<ext>"""
    import re
    base = re.sub(r'[^a-zA-Z0-9_-]+', '_', (label or '').strip()) or 'document'
    base = base[:40]
    ext = ''
    ft = (file_type or '').strip().lower()
    if ft == 'pdf':
        ext = '.pdf'
    elif ft == 'word':
        ext = '.docx'
    elif ft == 'excel':
        ext = '.xlsx'
    elif ft == 'image':
        ext = '.png'
    return f"intern-{intern_id}-{doc_id or 'tmp'}-{base}{ext}"


def safe_filename(prefix: str, original: str) -> str:
    """Generate a safe, unique upload filename while preserving the original extension."""
    import uuid
    base = secure_filename(original)
    ext = ''
    if '.' in original:
        ext = '.' + original.rsplit('.', 1)[1].lower()
        if len(ext) > 10 or not ext[1:].isalnum():
            ext = ''
    if not ext and base and '.' in base:
        ext = '.' + base.rsplit('.', 1)[1].lower()
    
    if not base:
        base = f"{prefix}_{uuid.uuid4().hex}"
    else:
        base = base.rsplit('.', 1)[0]
    
    # ALWAYS append a UUID to prevent race conditions and overwrites
    name = f"{base}_{uuid.uuid4().hex[:8]}{ext}"
    return name


def _uniquify_title(intern_id, title):
    if not title:
        return title
    exact = DocumentLifecycle.query.filter(
        DocumentLifecycle.intern_id == intern_id,
        DocumentLifecycle.custom_title == title
    ).first()
    if not exact:
        return title
    n = 1
    while DocumentLifecycle.query.filter(
        DocumentLifecycle.intern_id == intern_id,
        DocumentLifecycle.custom_title == f'{title} ({n})'
    ).first():
        n += 1
    return f'{title} ({n})'


def _parse_date(val, us_format=False):
    """Safely convert a date string to a Python date object for SQLite.

    us_format=True prefers US order (mm/dd/yyyy) used by Google Forms.
    """
    if not val:
        return None
    if isinstance(val, date):
        return val
    if not isinstance(val, str):
        return None
    val = val.strip()
    if us_format:
        formats = ('%m/%d/%Y', '%m-%d-%Y', '%m/%d/%y', '%Y-%m-%d', '%d/%m/%Y')
    else:
        formats = ('%Y-%m-%d', '%d/%m/%Y', '%m/%d/%Y', '%d-%m-%Y', '%Y/%m/%d')
    for fmt in formats:
        try:
            return datetime.strptime(val, fmt).date()
        except ValueError:
            continue
    return None


def log_action(user, action):
    try:
        new_log = SystemLog(user=user, action=action)
        db.session.add(new_log)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"Failed to log action: {e}")


# --- ROLE-BASED ACCESS CONTROL (RBAC) ---
# Granular permission matrix stored as JSON in User.permissions.
# Modules mirror the UI groups; actions: view / add / edit / delete / approve.
PERMISSION_ACTIONS = ('view', 'add', 'edit', 'delete', 'approve')

DEFAULT_PERMISSIONS = {
    # إدارة المتدربين والطلبات
    "interns":          {"view": True,  "add": True,  "edit": True,  "delete": False, "approve": True},
    "forms":            {"view": True,  "add": True,  "edit": True,  "delete": False, "approve": False},
    "approve_interns":  {"view": True,  "add": True,  "edit": True,  "delete": False, "approve": True},
    "assign_encadrant": {"view": True,  "add": True,  "edit": True,  "delete": False, "approve": False},
    # الوثائق والمستندات
    "vault":            {"view": True,  "add": True,  "edit": True,  "delete": False, "approve": False},
    "doc_templates":    {"view": True,  "add": True,  "edit": True,  "delete": False, "approve": False},
    "doc_reupload":     {"view": True,  "add": True,  "edit": True,  "delete": False, "approve": True},
    "attestation":      {"view": True,  "add": False, "edit": False, "delete": False, "approve": True},
    # المتابعة والتقييم
    "attendance":       {"view": True,  "add": True,  "edit": True,  "delete": False, "approve": False},
    "evaluate_interns": {"view": True,  "add": True,  "edit": True,  "delete": False, "approve": False},
    "tasks_projects":   {"view": True,  "add": True,  "edit": True,  "delete": False, "approve": False},
    # النظام والإعدادات (مقيدة افتراضياً للمشرفين)
    "roles":            {"view": False, "add": False, "edit": False, "delete": False, "approve": False},
    "system_settings":  {"view": False, "add": False, "edit": False, "delete": False, "approve": False},
    "activity_logs":    {"view": False, "add": False, "edit": False, "delete": False, "approve": False},
    "statistics":       {"view": True,  "add": False, "edit": False, "delete": False, "approve": False},
}

def merge_permissions(permissions_json):
    """Merge stored permissions JSON over the defaults so every module/action has a value."""
    stored = {}
    if permissions_json:
        try:
            stored = json.loads(permissions_json)
        except Exception:
            stored = {}
    if not isinstance(stored, dict):
        stored = {}
    merged = {}
    for mod, acts in DEFAULT_PERMISSIONS.items():
        mod_map = stored.get(mod, {}) if isinstance(stored.get(mod), dict) else {}
        merged[mod] = {act: bool(mod_map.get(act, default)) for act, default in acts.items()}
    return merged

def user_permissions(user):
    """Return the effective permission map for a user. None = full access (Admin)."""
    if user.role == 'Admin':
        return None
    return merge_permissions(user.permissions)

def has_permission(user, module, action):
    perms = user_permissions(user)
    if perms is None:
        return True
    return bool(perms.get(module, {}).get(action))

def check_permission(module, action):
    """Decorator enforcing real-time RBAC against the DB. Admin always bypasses; Interns are denied."""
    def decorator(fn):
        @functools.wraps(fn)
        def wrapper(*args, **kwargs):
            role = get_jwt().get('role')
            if role == 'Admin':
                return fn(*args, **kwargs)
            if role != 'Manager':
                return jsonify({"msg": "Unauthorized"}), 403
            user = db.session.get(User, get_jwt_identity())
            if not user or not has_permission(user, module, action):
                return jsonify({"msg": "Access Denied: لا تملك صلاحية لهذا الإجراء"}), 403
            return fn(*args, **kwargs)
        return wrapper
    return decorator

# Old smtp logic removed. We now use email_service.py for Gmail integration.
def send_email(to_email: str, subject: str, body_html: str):
    email_service.send_email(to_email, subject, body_html)


def init_db():
    with app.app_context():
        db.create_all()
        # Auto-add newer columns to existing DBs
        try:
            cols = [c['name'] for c in db.inspect(db.engine).get_columns('interns')]
            if 'evaluation' not in cols:
                db.session.execute(db.text('ALTER TABLE interns ADD COLUMN evaluation TEXT'))
                db.session.commit()
        except Exception as e:
            db.session.rollback()
            print(f"Migration check failed: {e}")
        # Auto-add specialty column
        try:
            cols = [c['name'] for c in db.inspect(db.engine).get_columns('interns')]
            if 'specialty' not in cols:
                db.session.execute(db.text('ALTER TABLE interns ADD COLUMN specialty VARCHAR(150)'))
                db.session.commit()
        except Exception as e:
            db.session.rollback()
            print(f"Specialty migration failed: {e}")
        # Auto-add username column to users table
        try:
            ucols = [c['name'] for c in db.inspect(db.engine).get_columns('users')]
            if 'username' not in ucols:
                db.session.execute(db.text('ALTER TABLE users ADD COLUMN username VARCHAR(80)'))
                db.session.commit()
        except Exception as e:
            db.session.rollback()
            print(f"User migration failed: {e}")
        # Populate username from email for existing users
        try:
            for u in User.query.filter_by(username=None).all():
                if u.email:
                    u.username = u.email.split('@')[0]
                else:
                    u.username = f"user_{u.id}"
            db.session.commit()
        except Exception:
            db.session.rollback()
        # Fix existing users where username equals email
        try:
            for u in User.query.all():
                if u.username and u.email and u.username == u.email:
                    u.username = u.email.split('@')[0]
            db.session.commit()
        except Exception:
            db.session.rollback()
        # Add requires_return column if missing
        try:
            with db.engine.connect() as conn:
                from sqlalchemy import text
                conn.execute(text("ALTER TABLE document_lifecycle ADD COLUMN requires_return BOOLEAN DEFAULT 0"))
                conn.commit()
        except Exception:
            db.session.rollback()
        try:
            with db.engine.connect() as conn:
                from sqlalchemy import text
                conn.execute(text("ALTER TABLE document_lifecycle ADD COLUMN returned_file_path VARCHAR(255)"))
                conn.commit()
        except Exception:
            db.session.rollback()
        # Add file_type / action_type / requested_by columns (dynamic document requests)
        try:
            dl_cols = [c['name'] for c in db.inspect(db.engine).get_columns('document_lifecycle')]
            with db.engine.connect() as conn:
                from sqlalchemy import text as sql_text
                if 'file_type' not in dl_cols:
                    conn.execute(sql_text("ALTER TABLE document_lifecycle ADD COLUMN file_type VARCHAR(20) DEFAULT 'pdf'"))
                if 'action_type' not in dl_cols:
                    conn.execute(sql_text("ALTER TABLE document_lifecycle ADD COLUMN action_type VARCHAR(20) DEFAULT 'view'"))
                if 'requested_by' not in dl_cols:
                    conn.execute(sql_text("ALTER TABLE document_lifecycle ADD COLUMN requested_by VARCHAR(20) DEFAULT 'ADMIN'"))
                if 'revision_requested_by' not in dl_cols:
                    conn.execute(sql_text("ALTER TABLE document_lifecycle ADD COLUMN revision_requested_by VARCHAR(20)"))
                conn.commit()
        except Exception as e:
            print(f"document_lifecycle migration: {e}")
        # Migrate Form table — add new columns if missing
        try:
            form_cols = [c['name'] for c in db.inspect(db.engine).get_columns('forms')]
            with db.engine.connect() as conn:
                from sqlalchemy import text as sql_text
                if 'title' not in form_cols:
                    conn.execute(sql_text('ALTER TABLE forms ADD COLUMN title VARCHAR(200)'))
                if 'slug' not in form_cols:
                    conn.execute(sql_text('ALTER TABLE forms ADD COLUMN slug VARCHAR(50)'))
                if 'is_active' not in form_cols:
                    conn.execute(sql_text('ALTER TABLE forms ADD COLUMN is_active BOOLEAN DEFAULT 1'))
                if 'created_at' not in form_cols:
                    conn.execute(sql_text('ALTER TABLE forms ADD COLUMN created_at VARCHAR(50)'))
                conn.commit()
        except Exception as e:
            db.session.rollback()
            print(f"Form migration: {e}")

        # Migrate: add can_manage_documents column to users table
        try:
            user_cols = [c['name'] for c in db.inspect(db.engine).get_columns('users')]
            if 'can_manage_documents' not in user_cols:
                with db.engine.connect() as conn:
                    conn.execute(sql_text("ALTER TABLE users ADD COLUMN can_manage_documents BOOLEAN DEFAULT 0"))
                    conn.commit()
        except Exception as e:
            db.session.rollback()
            print(f"User migration failed: {e}")

        # Create default Admin user if none exists
        admin = User.query.filter_by(username='admin').first()
        if not admin:
            # Check if there's an admin by email for backward compat
            admin = User.query.filter_by(email='admin@mahkama.ma').first()
            if admin:
                admin.username = 'admin'
            else:
                hashed_pw = generate_password_hash('admin123')
                admin = User(username='admin', name='مدير النظام', password=hashed_pw, role='Admin')
                db.session.add(admin)
            db.session.commit()

        # DocumentRequest lifecycle_id
        try:
            req_cols = [c['name'] for c in db.inspect(db.engine).get_columns('document_requests')]
            with db.engine.connect() as conn:
                from sqlalchemy import text as sql_text
                if 'lifecycle_id' not in req_cols:
                    conn.execute(sql_text('ALTER TABLE document_requests ADD COLUMN lifecycle_id INTEGER'))
                conn.commit()
        except Exception as e:
            db.session.rollback()
            print(f"DocumentRequest migration failed: {e}")

        # DocumentLifecycle new columns
        try:
            dl_cols = [c['name'] for c in db.inspect(db.engine).get_columns('document_lifecycle')]
            with db.engine.connect() as conn:
                from sqlalchemy import text as sql_text
                if 'lifecycle_type' not in dl_cols:
                    conn.execute(sql_text("ALTER TABLE document_lifecycle ADD COLUMN lifecycle_type VARCHAR(20) DEFAULT 'SIGN'"))
                if 'form_schema' not in dl_cols:
                    conn.execute(sql_text('ALTER TABLE document_lifecycle ADD COLUMN form_schema TEXT'))
                if 'form_data' not in dl_cols:
                    conn.execute(sql_text('ALTER TABLE document_lifecycle ADD COLUMN form_data TEXT'))
                if 'assigned_to_intern_id' not in dl_cols:
                    conn.execute(sql_text('ALTER TABLE document_lifecycle ADD COLUMN assigned_to_intern_id INTEGER'))
                if 'signed_by' not in dl_cols:
                    conn.execute(sql_text('ALTER TABLE document_lifecycle ADD COLUMN signed_by VARCHAR(150)'))
                if 'correction_round' not in dl_cols:
                    conn.execute(sql_text('ALTER TABLE document_lifecycle ADD COLUMN correction_round INTEGER DEFAULT 0'))
                if 'parent_id' not in dl_cols:
                    conn.execute(sql_text('ALTER TABLE document_lifecycle ADD COLUMN parent_id INTEGER'))
                conn.commit()
        except Exception as e:
            db.session.rollback()
            print(f"DocumentLifecycle migration failed: {e}")

        # Add returned_files_history column
        try:
            dl_cols = [c['name'] for c in db.inspect(db.engine).get_columns('document_lifecycle')]
            with db.engine.connect() as conn:
                from sqlalchemy import text as sql_text
                if 'returned_files_history' not in dl_cols:
                    conn.execute(sql_text('ALTER TABLE document_lifecycle ADD COLUMN returned_files_history TEXT'))
                conn.commit()
        except Exception as e:
            db.session.rollback()
            print(f"returned_files_history migration failed: {e}")

        # Add source column to document_lifecycle + backfill TEMPLATE_VIEW
        try:
            dl_cols = [c['name'] for c in db.inspect(db.engine).get_columns('document_lifecycle')]
            with db.engine.connect() as conn:
                from sqlalchemy import text as sql_text
                if 'source' not in dl_cols:
                    conn.execute(sql_text("ALTER TABLE document_lifecycle ADD COLUMN source VARCHAR(20) DEFAULT 'ADMIN'"))
                conn.commit()
            # Backfill old template-created rows — action_type='view', custom_title matches
            # an active template label. Stamp them as TEMPLATE_VIEW (regardless of file_path).
            with app.app_context():
                active_titles = set()
                try:
                    for t in DocumentTemplate.query.filter_by(is_active=True).all():
                        active_titles.add(t.label)
                except Exception:
                    active_titles = set()
                if active_titles:
                    rows = DocumentLifecycle.query.filter(
                        DocumentLifecycle.action_type == 'view',
                        DocumentLifecycle.custom_title.in_(active_titles),
                        (DocumentLifecycle.source == 'ADMIN') | (DocumentLifecycle.source.is_(None))
                    ).all()
                    for r in rows:
                        r.source = 'TEMPLATE_VIEW'
                    if rows:
                        db.session.commit()
                        print(f"Backfilled {len(rows)} template-created docs with source=TEMPLATE_VIEW")
        except Exception as e:
            db.session.rollback()
            print(f"source column migration/backfill failed: {e}")

        # Create DocumentTemplate table
        try:
            db.session.execute(sql_text("""
                CREATE TABLE IF NOT EXISTS document_templates (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    label VARCHAR(150) NOT NULL,
                    file_type VARCHAR(20) DEFAULT 'pdf',
                    file_path VARCHAR(255),
                    is_active BOOLEAN DEFAULT 1,
                    created_at DATETIME
                )
            """))
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            print(f"DocumentTemplate migration: {e}")
        # Add file_path to existing document_templates if missing
        try:
            dt_cols = [c['name'] for c in db.inspect(db.engine).get_columns('document_templates')]
            with db.engine.connect() as conn:
                if 'file_path' not in dt_cols:
                    conn.execute(sql_text("ALTER TABLE document_templates ADD COLUMN file_path VARCHAR(255)"))
                    conn.commit()
        except Exception as e:
            db.session.rollback()
            print(f"document_templates file_path migration: {e}")

        # Date string to DATE migrations
        try:
            with db.engine.connect() as conn:
                from sqlalchemy import text as sql_text
                # SQLite ALTER TABLE doesn't support changing column types easily.
                # Since sqlite is loosely typed, we don't strictly need to alter the schema 
                # type in SQLite, but we can do a data normalization pass here if we wanted.
                # The python models will be updated to Date/DateTime.
                pass
        except Exception as e:
            db.session.rollback()

        # Migrate legacy sign/fill lifecycle docs to requires_return=True
        try:
            fixed = DocumentLifecycle.query.filter(
                DocumentLifecycle.action_type.in_(['sign', 'fill', 'sign_fill']),
                DocumentLifecycle.requires_return == False
            ).update(
                {'requires_return': True},
                synchronize_session='fetch'
            )
            db.session.commit()
            if fixed:
                print(f"Fixed {fixed} legacy sign/fill lifecycle docs (requires_return=True)")
        except Exception as e:
            db.session.rollback()
            print(f"Legacy sign/fill migration failed: {e}")

        # Migrate legacy sign/fill request slots to be visible to intern
        try:
            fixed2 = DocumentLifecycle.query.filter(
                DocumentLifecycle.action_type.in_(['sign', 'fill', 'sign_fill']),
                DocumentLifecycle.is_visible_to_intern == False
            ).update(
                {'is_visible_to_intern': True, 'requires_return': True},
                synchronize_session='fetch'
            )
            db.session.commit()
            if fixed2:
                print(f"Fixed {fixed2} legacy sign/fill lifecycle docs (is_visible_to_intern=True)")
        except Exception as e:
            db.session.rollback()
            print(f"Legacy sign/fill visibility migration failed: {e}")

        # Fix: revert AWAITING_RETURN docs that have no file_path back to MISSING
        try:
            fixed3 = DocumentLifecycle.query.filter(
                DocumentLifecycle.status == 'AWAITING_RETURN',
                DocumentLifecycle.file_path.is_(None)
            ).update(
                {'status': 'MISSING'},
                synchronize_session='fetch'
            )
            db.session.commit()
            if fixed3:
                print(f"Fixed {fixed3} AWAITING_RETURN docs with no file (reverted to MISSING)")
        except Exception as e:
            db.session.rollback()
            print(f"Status fixup migration failed: {e}")


# --- AUTHENTICATION ROUTES ---
@app.route('/api/login', methods=['POST'])
def login():
    data = request.json
    username = data.get('username', data.get('email', '')).strip()
    password = data.get('password')
    
    user = User.query.filter((User.username == username) | (User.email == username)).first()
    
    if user and check_password_hash(user.password, password):
        access_token = create_access_token(
            identity=str(user.id),
            additional_claims={
                'name': user.name, 
                'role': user.role, 
                'permissions': user.permissions,
                'can_manage_documents': user.can_manage_documents
            }
        )
        log_action(user.name, "قام بتسجيل الدخول إلى النظام")
        return jsonify(
            access_token=access_token, 
            user={'id': user.id, 'name': user.name, 'role': user.role, 'permissions': user.permissions, 'can_manage_documents': user.can_manage_documents}
        ), 200
        
    return jsonify({"msg": "اسم المستخدم أو كلمة المرور غير صحيحة"}), 401


@app.route('/api/auth/me', methods=['GET'])
@jwt_required()
def get_me():
    """Return fresh identity + effective permissions (bypasses stale JWT claims)."""
    user = db.session.get(User, get_jwt_identity())
    if not user:
        return jsonify({"msg": "User not found"}), 401
    perms = user_permissions(user)
    return jsonify({
        "id": user.id, "name": user.name, "username": user.username,
        "email": user.email, "role": user.role,
        "permissions": perms,
        "can_manage_documents": user.can_manage_documents
    }), 200


# --- USERS ROUTES (Admin Only) ---
@app.route('/api/users', methods=['GET'])
@jwt_required()
@check_permission('roles', 'view')
def get_users():
    users = User.query.all()
    return jsonify([{"id": u.id, "name": u.name, "username": u.username, "email": u.email, "role": u.role, "permissions": u.permissions, "can_manage_documents": u.can_manage_documents} for u in users])

@app.route('/api/users', methods=['POST'])
@jwt_required()
@check_permission('roles', 'add')
def add_user():
    current_user = get_jwt()
    data = request.json
    
    # Check if username exists
    username = (data.get('username') or '').strip()
    if not username:
        email = (data.get('email') or '').strip()
        username = email.split('@')[0] if email else f"user_{User.query.count() + 1}"
        base_username = username
        counter = 1
        while User.query.filter_by(username=username).first():
            username = f"{base_username}{counter}"
            counter += 1
    if User.query.filter_by(username=username).first():
        return jsonify({"msg": "اسم المستخدم موجود بالفعل"}), 400

    email = (data.get('email') or '').strip()
    if email and User.query.filter_by(email=email).first():
        return jsonify({"msg": "البريد الإلكتروني مستخدم بالفعل"}), 400

    hashed_pw = generate_password_hash(data.get('password', 'password123'))
    new_user = User(
        username=username,
        name=data.get('name'), 
        email=email, 
        password=hashed_pw, 
        role=data.get('role'), 
        permissions=data.get('permissions', ''),
        can_manage_documents=data.get('can_manage_documents', False) if data.get('role') == 'Manager' else False
    )
    
    db.session.add(new_user)
    db.session.commit()
    
    log_action(current_user.get('name', 'Admin'), f"قام بإضافة مستخدم جديد: {new_user.name}")
    
    return jsonify({"success": True, "id": new_user.id})

@app.route('/api/users/<int:user_id>', methods=['PUT'])
@jwt_required()
@check_permission('roles', 'edit')
def update_user(user_id):
    current_user = get_jwt()
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"msg": "User not found"}), 404
        
    data = request.json
    user.name = data.get('name', user.name)
    user.username = data.get('username', user.username)
    user.email = data.get('email', user.email)
    user.role = data.get('role', user.role)
    if 'can_manage_documents' in data and user.role == 'Manager':
        user.can_manage_documents = data.get('can_manage_documents')
    if 'permissions' in data:
        import json
        user.permissions = data.get('permissions')
        
    if data.get('password'):
        user.password = generate_password_hash(data.get('password'))
        
    db.session.commit()
    log_action(current_user.get('name', 'Admin'), f"قام بتحديث بيانات المستخدم: {user.name}")
    return jsonify({"success": True})

@app.route('/api/users/password', methods=['PUT'])
@jwt_required()
def change_password():
    current_user = get_jwt()
    user_id = current_user.get('sub')
    data = request.json
    
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"msg": "User not found"}), 404
        
    old_password = data.get('old_password')
    new_password = data.get('new_password')
    
    if not old_password or not new_password:
        return jsonify({"msg": "Old and new password required"}), 400
        
    from werkzeug.security import check_password_hash
    if not check_password_hash(user.password, old_password):
        return jsonify({"msg": "كلمة المرور القديمة غير صحيحة"}), 400
        
    hashed_pw = generate_password_hash(new_password)
    user.password = hashed_pw
    db.session.commit()
    
    log_action(user.name, "قام بتغيير كلمة المرور الخاصة به")
    return jsonify({"success": True})

@app.route('/api/users/<int:user_id>', methods=['DELETE'])
@jwt_required()
@check_permission('roles', 'delete')
def delete_user(user_id):
    current_user = get_jwt()
    user = db.session.get(User, user_id)
    if user:
        name_deleted = user.name
        db.session.delete(user)
        db.session.commit()
        log_action(current_user.get('name', 'Admin'), f"قام بحذف المستخدم: {name_deleted}")
    return jsonify({"success": True})

@app.route('/api/users/<int:user_id>/reset-password', methods=['POST'])
@jwt_required()
@check_permission('roles', 'edit')
def reset_user_password(user_id):
    current_user = get_jwt()
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"msg": "User not found"}), 404
    from werkzeug.security import generate_password_hash
    user.password = generate_password_hash('password123')
    db.session.commit()
    log_action(current_user.get('name', 'Admin'), f"قام بإعادة تعيين كلمة مرور المستخدم: {user.name}")
    return jsonify({"success": True, "msg": f"تم إعادة تعيين كلمة مرور {user.name} إلى password123"})


# --- INTERN ROUTES ---
@app.route('/api/interns', methods=['GET'])
@jwt_required()
@check_permission('interns', 'view')
def get_interns():
    from sqlalchemy import or_, and_
    interns = Intern.query.all()
    activity_statuses = {'PENDING_REVIEW', 'RETURNED', 'AWAITING_ADMIN'}
    return jsonify([{
        "id": i.id, 
        "name": i.name, 
        "email": i.email, 
        "encadrant": i.encadrant, 
        "name_fr": i.name_fr,
        "national_id": i.national_id,
        "status": i.status,
        "photo_path": i.photo_path,
        "start_date": i.start_date,
        "end_date": i.end_date,
        "department": i.department,
        "source": i.source,
        "has_final_report": DocumentLifecycle.query.filter_by(intern_id=i.id, doc_type='FINAL_REPORT').first() is not None,
        "has_pending_activity": DocumentLifecycle.query.filter(
            DocumentLifecycle.intern_id == i.id,
            or_(
                DocumentLifecycle.status == 'AWAITING_ADMIN',
                and_(
                    DocumentLifecycle.status.in_(activity_statuses),
                    DocumentLifecycle.uploaded_by != 'ADMIN'
                ),
                and_(
                    DocumentLifecycle.status == 'REVISION_REQUESTED',
                    DocumentLifecycle.revision_requested_by == 'INTERN'
                )
            )
        ).first() is not None,
        "has_missing_documents": DocumentLifecycle.query.filter(
            DocumentLifecycle.intern_id == i.id,
            DocumentLifecycle.status.in_(['MISSING', 'AWAITING_INTERN', 'AWAITING_RETURN'])
        ).first() is not None
    } for i in interns])

def create_user_for_intern(intern_obj):
    if intern_obj.email:
        existing_user = User.query.filter_by(email=intern_obj.email).first()
        if not existing_user:
            from werkzeug.security import generate_password_hash
            hashed_pw = generate_password_hash('password123')
            # Ensure unique username
            base_username = intern_obj.email.split('@')[0]
            username = base_username
            counter = 1
            while User.query.filter_by(username=username).first():
                username = f"{base_username}{counter}"
                counter += 1
                
            new_user = User(
                username=username,
                name=intern_obj.name or 'متدرب',
                email=intern_obj.email,
                password=hashed_pw,
                role='Intern',
                permissions=''
            )
            db.session.add(new_user)
            db.session.commit()

@app.route('/api/interns/<int:intern_id>', methods=['GET'])
@jwt_required()
@check_permission('interns', 'view')
def get_intern(intern_id):
    intern = db.session.get(Intern, intern_id)
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404
        
    import json
    docs = {}
    if intern.documents:
        try:
            docs = json.loads(intern.documents)
        except Exception as e: print(f"Silent Error Caught: {e}")

    evaluation = None
    if intern.evaluation:
        try:
            evaluation = json.loads(intern.evaluation)
        except Exception as e: print(f"Silent Error Caught: {e}")
            
    return jsonify({
        "id": intern.id, 
        "name": intern.name,
        "name_fr": intern.name_fr,
        "email": intern.email, 
        "national_id": intern.national_id,
        "department": intern.department, 
        "encadrant": intern.encadrant,
        "status": intern.status,
        "photo_path": intern.photo_path,
        "phone": intern.phone,
        "start_date": intern.start_date,
        "end_date": intern.end_date,
        "date_of_birth": intern.date_of_birth,
        "university": intern.university,
        "specialty": intern.specialty,
        "address": intern.address,
        "documents": docs,
        "evaluation": evaluation
    })

@app.route('/api/interns', methods=['POST'])
@jwt_required()
@check_permission('interns', 'add')
def add_intern():
    data = request.json
    new_intern = create_intern_record(data, 'إضافة يدوية')
    db.session.commit()

    current_user = get_jwt()
    user_name = current_user.get('name') if current_user else 'Unknown'
    log_action(user_name, f"قام بإضافة متدرب جديد: {new_intern.name}")

    return jsonify({"success": True, "id": new_intern.id})

@app.route('/api/interns/<int:intern_id>', methods=['PUT'])
@jwt_required()
@check_permission('interns', 'edit')
def update_intern(intern_id):
    intern = db.session.get(Intern, intern_id)
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404
        
    data = request.json
    import json
    
    current_user = get_jwt()
    if 'encadrant' in data and data['encadrant'] != intern.encadrant:
        if not has_permission(db.session.get(User, get_jwt_identity()), 'assign_encadrant', 'edit'):
            return jsonify({"msg": "Access Denied: لا تملك صلاحية تعيين المؤطر/المشرف"}), 403
    if 'status' in data and data['status'] != intern.status:
        if current_user.get('role') != 'Admin':
            return jsonify({"msg": "Unauthorized: Only Admins can change intern status"}), 403
        # Auto-create lifecycle entries when activated
        if data['status'] == 'نشط':
            if intern.email:
                try:
                    email_service.send_accepted_email(intern.email, intern.name)
                except Exception as e:
                    print("Error sending acceptance email:", e)
            templates = DocumentTemplate.query.filter_by(is_active=True).all()
            now = datetime.now(timezone.utc)
            existing_titles = {d.custom_title for d in DocumentLifecycle.query.filter_by(intern_id=intern.id).all()}
            for t in templates:
                if t.label not in existing_titles:
                    record = DocumentLifecycle(
                        intern_id=intern.id, doc_type='OTHER', status='MISSING',
                        uploaded_by='ADMIN', is_visible_to_intern=True,
                        custom_title=t.label, action_type='view',
                        created_at=now, updated_at=now,
                        source='TEMPLATE_VIEW'
                    )
                    db.session.add(record)
            if templates:
                db.session.commit()
        elif data['status'] == 'مرفوض':
            if intern.email:
                try:
                    email_service.send_rejected_email(intern.email, intern.name)
                except Exception as e:
                    print("Error sending rejection email:", e)
            
    intern.name = data.get('name', intern.name)
    if 'name_fr' in data:
        intern.name_fr = data.get('name_fr')
    intern.email = data.get('email', intern.email)
    intern.national_id = data.get('national_id', intern.national_id)
    intern.department = data.get('department', intern.department)
    if 'encadrant' in data:
        intern.encadrant = data.get('encadrant')
    intern.phone = data.get('phone', intern.phone)
    intern.start_date = _parse_date(data.get('start_date')) if 'start_date' in data else intern.start_date
    intern.end_date = _parse_date(data.get('end_date')) if 'end_date' in data else intern.end_date
    intern.date_of_birth = _parse_date(data.get('date_of_birth')) if 'date_of_birth' in data else intern.date_of_birth
    intern.university = data.get('university', intern.university)
    intern.specialty = data.get('specialty', intern.specialty)
    intern.address = data.get('address', intern.address)
    if 'status' in data:
        intern.status = data.get('status')
    
    if 'photo_path' in data:
        intern.photo_path = data.get('photo_path')
    if 'documents' in data:
        intern.documents = json.dumps(data.get('documents'))
        
    db.session.commit()
    
    current_user = get_jwt()
    user_name = current_user.get('name') if current_user else 'Unknown'
    log_action(user_name, f"قام بتعديل بيانات المتدرب: {intern.name}")
    
    return jsonify({"success": True})

@app.route('/api/interns/<int:intern_id>/evaluation', methods=['POST'])
@jwt_required()
@check_permission('evaluate_interns', 'add')
def save_evaluation(intern_id):
    intern = db.session.get(Intern, intern_id)
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404

    import json
    data = request.json or {}
    current_user = get_jwt()
    user_name = current_user.get('name') if current_user else 'Unknown'

    existing = {}
    if intern.evaluation:
        try: existing = json.loads(intern.evaluation)
        except Exception as e: print(f"Silent Error Caught: {e}")

    evaluation = {
        "training_location": data.get('training_location', existing.get('training_location', '')),
        "period_from": data.get('period_from', existing.get('period_from', '')),
        "period_to": data.get('period_to', existing.get('period_to', '')),
        "rotations": data.get('rotations', existing.get('rotations', [])),
        "criteria": data.get('criteria', existing.get('criteria', {})),
        "comments": data.get('comments', existing.get('comments', '')),
        "signed_file_path": existing.get('signed_file_path'),
        "evaluator": user_name,
        "date": datetime.now().strftime('%Y-%m-%d %H:%M')
    }
    intern.evaluation = json.dumps(evaluation, ensure_ascii=False)
    db.session.commit()

    log_action(user_name, f"قام بتقييم المتدرب: {intern.name}")
    return jsonify({"success": True, "evaluation": evaluation})

@app.route('/api/interns/<int:intern_id>/evaluation/signed-upload', methods=['POST'])
@jwt_required()
@check_permission('evaluate_interns', 'edit')
def upload_signed_evaluation(intern_id):
    intern = db.session.get(Intern, intern_id)
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404
    current_user = get_jwt()

    if 'file' not in request.files:
        return jsonify({"msg": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"msg": "No selected file"}), 400
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({"msg": "PDF files only"}), 400
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > 15 * 1024 * 1024:
        return jsonify({"msg": "File too large (max 15MB)"}), 400

    filename = safe_filename('eval_signed', file.filename)
    if not filename.lower().endswith('.pdf'):
        filename += '.pdf'
    file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
    file_url = f"/api/uploads/{filename}"

    import json
    existing = {}
    if intern.evaluation:
        try: existing = json.loads(intern.evaluation)
        except Exception as e: print(f"Silent Error Caught: {e}")
    existing['signed_file_path'] = file_url
    intern.evaluation = json.dumps(existing, ensure_ascii=False)
    db.session.commit()

    log_action(current_user.get('name'), f"رفع نسخة موقعة من تقييم المتدرب {intern.name}")
    return jsonify({"success": True, "signed_file_path": file_url}), 200

@app.route('/api/interns/<int:intern_id>/attendance', methods=['GET'])
@jwt_required()
@check_permission('attendance', 'view')
def get_attendance(intern_id):
    records = Attendance.query.filter_by(intern_id=intern_id).order_by(Attendance.date.desc()).all()
    return jsonify([{"id": r.id, "date": r.date, "status": r.status} for r in records])

@app.route('/api/interns/<int:intern_id>/attendance', methods=['POST'])
@jwt_required()
@check_permission('attendance', 'add')
def mark_attendance(intern_id):
    data = request.json
    date = data.get('date')
    status = data.get('status')
    
    if not date or not status:
        return jsonify({"msg": "Date and status required"}), 400
        
    # Check if record exists for this date
    record = Attendance.query.filter_by(intern_id=intern_id, date=date).first()
    if record:
        record.status = status
    else:
        record = Attendance(intern_id=intern_id, date=date, status=status)
        db.session.add(record)
        
    db.session.commit()
    return jsonify({"success": True, "id": record.id})

@app.route('/api/forms/sync-microsoft', methods=['POST'])
@jwt_required()
@check_permission('forms', 'add')
def sync_microsoft_forms():
    import microsoft_excel_service
    settings = email_service.get_settings()
    excel_link = settings.get('microsoft_excel_link')
    if not excel_link:
        return jsonify({"success": False, "msg": "يرجى تعيين رابط Microsoft Excel في الإعدادات أولاً"})
        
    res = microsoft_excel_service.fetch_microsoft_excel_responses(excel_link)
    if not res['success']:
        return jsonify(res), 400
        
    rows = res['data']
    added_count = 0
    
    # We use 'Microsoft Excel' as the form title to distinguish
    form_title = "Microsoft Excel"
    
    for row in rows:
        # Generate a stable hash or identifier if possible, but Excel rows usually lack timestamps unless the user added them.
        # Try to find an email or name to deduplicate
        email = None
        for k, v in row.items():
            if 'email' in k.lower() or 'بريد' in k:
                email = str(v)
                break
                
        # Simple deduplication based on exact data match
        existing = FormSubmission.query.all()
        is_duplicate = False
        for ex in existing:
            # If the JSON data is exactly the same, skip
            if ex.submitted_data == row:
                is_duplicate = True
                break
                
        if is_duplicate:
            continue
            
        default_form = Form.query.first()
        new_sub = FormSubmission(
            form_id=default_form.id if default_form else 1,
            submitted_data=row,
            status='pending' # Will be approved instantly
        )
        db.session.add(new_sub)
        db.session.flush() # So it gets an ID
        
        # AUTO APPROVE IT
        new_intern = _process_submission_to_intern(new_sub)
        log_action(get_jwt_identity(), f"تم المزامنة التلقائية للطلب #{new_sub.id} كمتدرب #{new_intern.id}")
        
        added_count += 1
        
    db.session.commit()
    return jsonify({"success": True, "added": added_count})

@app.route('/api/forms/generate', methods=['POST'])
@jwt_required()
@check_permission('forms', 'add')
def generate_google_form_endpoint():
    data = request.json
    title = data.get("title", "نموذج جديد")
    fields = data.get("fields", [])
    
    settings = email_service.get_settings()
    client_id = settings.get("google_client_id")
    client_secret = settings.get("google_client_secret")
    owner_email = settings.get("gmail_address", "")
    
    if not client_id or not client_secret:
        return jsonify({"success": False, "msg": "يرجى إضافة Client ID و Client Secret في الإعدادات."}), 400
        
    res = google_forms_api.generate_google_form(title, fields, owner_email)
    if not res.get("success"):
        if res.get("needs_auth"):
            return jsonify(res), 401
        return jsonify(res), 400
        
    return jsonify(res)

@app.route('/api/oauth/google/url', methods=['GET'])
def get_google_auth_url():
    settings = email_service.get_settings()
    client_id = settings.get("google_client_id")
    client_secret = settings.get("google_client_secret")
    if not client_id or not client_secret:
        return jsonify({"success": False, "msg": "إعدادات OAuth مفقودة"}), 400
        
    try:
        url = google_forms_api.get_auth_url(client_id, client_secret)
        return jsonify({"success": True, "url": url})
    except Exception as e:
        return jsonify({"success": False, "msg": str(e)}), 500

@app.route('/oauth/callback', methods=['GET'])
def oauth_callback():
    from flask import redirect
    code = request.args.get('code')
    if not code:
        return "خطأ: لم يتم تلقي كود التحقق", 400
        
    settings = email_service.get_settings()
    client_id = settings.get("google_client_id")
    client_secret = settings.get("google_client_secret")
    
    try:
        google_forms_api.save_token_from_code(code, client_id, client_secret)
        return """
        <html>
            <head><meta charset="utf-8"><title>تم تسجيل الدخول</title></head>
            <body style="font-family: Arial, sans-serif; text-align: center; padding-top: 50px;">
                <h1 style="color: #4CAF50;">✅ تم تسجيل الدخول إلى حساب جوجل بنجاح!</h1>
                <p style="font-size: 18px;">يمكنك إغلاق هذه النافذة والعودة إلى تطبيق إدارة المتدربين لإنشاء النماذج.</p>
                <script>
                    setTimeout(() => window.close(), 3000);
                </script>
            </body>
        </html>
        """
    except Exception as e:
        return f"<html><body><h1 style='color:red'>حدث خطأ: {str(e)}</h1></body></html>", 500

@app.route('/api/attendance/by-date', methods=['GET'])
@jwt_required()
@check_permission('attendance', 'view')
def get_attendance_by_date():
    from datetime import date as dt_date
    date_str = request.args.get('date')
    if not date_str:
        date_str = dt_date.today().isoformat()
    records = Attendance.query.filter_by(date=date_str).all()
    # return a dict mapping intern_id to status
    att_dict = {r.intern_id: r.status for r in records}
    return jsonify(att_dict)

import io
from flask import send_file


@app.route('/api/attendance/generate-daily-record', methods=['GET'])
@jwt_required()
@check_permission('attendance', 'view')
def generate_daily_attendance_record():
    date_str = request.args.get('date')
    if not date_str:
        return jsonify({"msg": "Date is required"}), 400
        
    try:
        from datetime import datetime
        dt = datetime.strptime(date_str, '%Y-%m-%d')
    except ValueError:
        return jsonify({"msg": "Invalid date format"}), 400

    try:
        import docx
    except ImportError:
        return jsonify({"msg": "python-docx library not installed"}), 500

    template_path = os.path.join(os.path.dirname(__file__), '..', 'stagaires.docx')
    if not os.path.exists(template_path):
        return jsonify({"msg": "Template file stagaires.docx not found"}), 404

    # Get active interns whose dates overlap with the selected date
    all_interns = Intern.query.filter_by(status='نشط').all()
    active_interns = []
    for intern in all_interns:
        if intern.start_date:
            try:
                start_dt = datetime.strptime(intern.start_date, '%Y-%m-%d')
                if dt < start_dt:
                    continue
            except ValueError:
                pass
        if intern.end_date:
            try:
                end_dt = datetime.strptime(intern.end_date, '%Y-%m-%d')
                if dt > end_dt:
                    continue
            except ValueError:
                pass
        active_interns.append(intern)

    doc = docx.Document(template_path)
    
    # Find the paragraph with the date and update it while preserving style
    date_formatted = dt.strftime('%d/%m/%Y')
    for p in doc.paragraphs:
        if "ورقة الحضور بتاريخ" in p.text:
            new_text = p.text.replace(".../.../2026", date_formatted)
            if p.runs:
                p.runs[0].text = new_text
                for r in p.runs[1:]:
                    r.text = ""
            break
        
    if len(doc.tables) > 0:
        table = doc.tables[0]
        # Keep header, remove other rows
        for _ in range(len(table.rows) - 1):
            row = table.rows[1]
            row._element.getparent().remove(row._element)
            
        from docx.shared import Pt
        # Add new rows for active interns
        for intern in active_interns:
            row_cells = table.add_row().cells
            # Setting text with specific font and size
            p = row_cells[0].paragraphs[0]
            run = p.add_run(intern.name)
            run.font.name = 'Samir_Khouaja_Maghribi'
            run.font.size = Pt(20)
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].text = ""
            row_cells[2].text = ""

    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"سجل_الحضور_اليومي_{date_str}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/api/interns/<int:intern_id>/attestation', methods=['GET'])
@jwt_required()
@check_permission('attestation', 'approve')
def generate_attestation(intern_id):
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import bidi.algorithm as bidi
        import arabic_reshaper
    except ImportError:
        return jsonify({"msg": "PDF generation library not installed"}), 500
        
    intern = db.session.get(Intern, intern_id)
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404
        
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    
    # Try to load a font that supports Arabic if available, else fallback
    # For a real production app, we would bundle a font like Amiri or Tahoma.
    # We will use default Helvetica for French, and try to draw Arabic if possible.
    
    c.setFont("Helvetica-Bold", 24)
    c.drawCentredString(width/2.0, height - 100, "ATTESTATION DE STAGE")
    
    c.setFont("Helvetica", 14)
    text = f"Nous soussignes, certifions que Monsieur/Madame:"
    c.drawString(50, height - 200, text)
    
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 230, f"{intern.name_fr or '__________________'}")
    
    c.setFont("Helvetica", 14)
    c.drawString(50, height - 280, f"A effectue un stage au sein de notre departement:")
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, height - 310, f"{intern.department or '__________________'}")
    
    c.setFont("Helvetica", 14)
    c.drawString(50, height - 360, f"Encadre(e) par:")
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, height - 390, f"{intern.encadrant or '__________________'}")
    
    c.setFont("Helvetica", 14)
    c.drawString(50, height - 440, f"Période: du {intern.start_date or '___'} au {intern.end_date or '___'}")
    
    c.drawString(50, height - 500, "Cette attestation est delivree pour servir et valoir ce que de droit.")
    
    c.drawString(width - 200, 150, "Signature et Cachet:")
    
    c.showPage()
    c.save()
    
    buffer.seek(0)
    filename = f"Attestation_{intern.id}.pdf"
    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/pdf')

@app.route('/api/interns/<int:intern_id>', methods=['DELETE'])
@jwt_required()
@check_permission('interns', 'delete')
def delete_intern(intern_id):
    intern = db.session.get(Intern, intern_id)
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404
        
    name = intern.name
    DocumentLifecycle.query.filter_by(intern_id=intern_id).delete()
    Notification.query.filter_by(intern_id=intern_id).delete()
    Attendance.query.filter_by(intern_id=intern_id).delete()
    if intern.email:
        User.query.filter_by(email=intern.email, role='Intern').delete()
        if not BlacklistedIntern.query.filter_by(identifier=intern.email).first():
            db.session.add(BlacklistedIntern(identifier=intern.email))
    if intern.phone and not BlacklistedIntern.query.filter_by(identifier=intern.phone).first():
        db.session.add(BlacklistedIntern(identifier=intern.phone))
    db.session.delete(intern)
    db.session.commit()
    
    current_user = get_jwt()
    user_name = current_user.get('name') if current_user else 'Unknown'
    log_action(user_name, f"قام بحذف المتدرب: {name}")
    
    return jsonify({"success": True})


# --- DOCUMENT VAULT ROUTES (separate from intern profile documents) ---
VAULT_FOLDER = os.path.join(os.path.dirname(__file__), 'vault')
os.makedirs(VAULT_FOLDER, exist_ok=True)

@app.route('/api/vault', methods=['GET'])
@jwt_required()
@check_permission('vault', 'view')
def list_vault_documents():
    files = []
    if os.path.exists(VAULT_FOLDER):
        for filename in os.listdir(VAULT_FOLDER):
            filepath = os.path.join(VAULT_FOLDER, filename)
            if os.path.isfile(filepath):
                files.append({
                    "name": filename,
                    "size": os.path.getsize(filepath)
                })
    return jsonify(files)

@app.route('/api/vault', methods=['POST'])
@jwt_required()
@check_permission('vault', 'add')
def upload_vault_document():
    if 'file' not in request.files:
        return jsonify({"msg": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"msg": "No selected file"}), 400
        
    if file:
        if not (file.filename.lower().endswith('.pdf') or file.filename.lower().endswith('.docx') or file.filename.lower().endswith('.doc')):
            return jsonify({"msg": "عذراً، يُسمح فقط برفع ملفات PDF و Word"}), 400
            
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
        if size > 5 * 1024 * 1024:
            return jsonify({"msg": "عذراً، حجم الملف يجب أن لا يتجاوز 5 ميجابايت"}), 400
            
        custom_name = request.form.get('custom_name')
        if custom_name:
            filename = custom_name.replace('/', '_').replace('\\', '_')
            
            # Determine correct extension based on original file
            ext = '.pdf'
            if file.filename.lower().endswith('.docx'):
                ext = '.docx'
            elif file.filename.lower().endswith('.doc'):
                ext = '.doc'
                
            if not (filename.lower().endswith('.pdf') or filename.lower().endswith('.docx') or filename.lower().endswith('.doc')):
                filename += ext
        else:
            filename = secure_filename(file.filename)
            if not filename:
                filename = "vault_" + file.filename.replace('/', '_').replace('\\', '_')
        file.save(os.path.join(VAULT_FOLDER, filename))
        
        current_user = get_jwt()
        user_name = current_user.get('name') if current_user else 'Unknown'
        log_action(user_name, f"قام برفع مستند للخزنة: {filename}")
        
        return jsonify({"msg": "تم رفع الملف بنجاح", "filename": filename}), 201

@app.route('/api/vault/<filename>', methods=['GET'])
def download_vault_document(filename):
    return send_from_directory(VAULT_FOLDER, filename)

@app.route('/api/vault/open/<filename>', methods=['GET'])
def open_vault_document(filename):
    safe_name = filename.replace('/', '_').replace('\\', '_')
    filepath = os.path.join(VAULT_FOLDER, safe_name)
    if os.path.exists(filepath):
        import os as _os
        if _os.name == 'nt':
            try:
                _os.startfile(filepath)
                return jsonify({"success": True})
            except OSError:
                return jsonify({"msg": "لم يتم العثور على برنامج لفتح ملفات Word (مثل Microsoft Word) على جهازك."}), 400
        # If not Windows, we just return the file as a fallback, 
        # or we could try xdg-open/open, but Windows is what matters here.
    return jsonify({"msg": "File not found"}), 404

@app.route('/api/vault/<filename>', methods=['DELETE'])
@jwt_required()
@check_permission('vault', 'delete')
def delete_vault_document(filename):
    filepath = os.path.join(VAULT_FOLDER, filename)
    if os.path.exists(filepath):
        os.remove(filepath)
        current_user = get_jwt()
        user_name = current_user.get('name') if current_user else 'Unknown'
        log_action(user_name, f"قام بحذف مستند من الخزنة: {filename}")
        return jsonify({"success": True})
    return jsonify({"msg": "File not found"}), 404


@app.route('/api/interns/<int:intern_id>/vault-attach', methods=['POST'])
@jwt_required()
@check_permission('doc_reupload', 'add')
def attach_vault_to_intern(intern_id):
    intern = db.session.get(Intern, intern_id)
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404
    data = request.json or {}
    vault_name = data.get('vault_name')
    doc_type = data.get('doc_type', 'OTHER')
    action_type = data.get('action_type', 'view')
    custom_title = data.get('custom_title', '').strip() or None
    requires_return = action_type in ('sign', 'fill', 'sign_fill')
    if not vault_name:
        return jsonify({"msg": "vault_name required"}), 400
    src = os.path.join(VAULT_FOLDER, vault_name)
    if not os.path.exists(src):
        return jsonify({"msg": "Vault file not found"}), 404
    import shutil
    dst_name = f"vault_{intern_id}_{vault_name}"
    dst = os.path.join(app.config['UPLOAD_FOLDER'], dst_name)
    shutil.copy2(src, dst)
    file_url = f"/api/uploads/{dst_name}"
    now = datetime.now(timezone.utc)
    action_labels = {'sign': 'توقيع', 'fill': 'تعبئة وإرجاع', 'sign_fill': 'توقيع وتعبئة وإرجاع'}
    action_label = action_labels.get(action_type, '')
    label = _uniquify_title(intern.id, custom_title or f'{vault_name} ({action_label})')
    record = DocumentLifecycle(
        intern_id=intern.id, doc_type=doc_type, file_path=file_url,
        uploaded_by='ADMIN', status='AWAITING_RETURN' if requires_return else 'APPROVED_AND_SIGNED',
        is_visible_to_intern=True, custom_title=label,
        requires_return=requires_return, action_type=action_type,
        created_at=now, updated_at=now
    )
    db.session.add(record)
    db.session.commit()
    notif_titles = {'sign': 'مستند يطلب التوقيع', 'fill': 'مستند يطلب التعبئة', 'sign_fill': 'مستند يطلب التوقيع والتعبئة', 'upload': 'مستند مطلوب رفعه', 'view': 'مستند جديد للمعاينة'}
    notif_bodies = {
        'sign': f'المستند "{label}" يتطلب منك توقيعه.',
        'fill': f'المستند "{label}" يتطلب منك تعبئته وإرجاعه.',
        'sign_fill': f'المستند "{label}" يتطلب منك توقيعه وتعبئته وإرجاعه.',
        'upload': f'مطلوب منك رفع المستند "{label}".',
        'view': f'تم إرفاق المستند "{label}" لمعاينته.',
    }
    notif_type = 'GENERIC' if action_type == 'view' else 'REASSIGNMENT'
    notif = Notification(intern_id=intern.id, type=notif_type,
        title=notif_titles.get(action_type, 'مستند جديد'),
        body=notif_bodies.get(action_type, 'يوجد مستند جديد.'),
        related_doc_id=record.id)
    db.session.add(notif)
    db.session.commit()
    current_user = get_jwt()
    log_action(current_user.get('name'), f"أضاف مستند من الخزنة ({vault_name}){' مع طلب التعبئة' if requires_return else ''} للمتدرب {intern.name}")
    return jsonify({"msg": "تمت الإضافة من الخزنة", "doc_id": record.id}), 200

# --- INTERN PROFILE DOCUMENT UPLOADS ---
@app.route('/api/documents', methods=['POST'])
@jwt_required()
def upload_document():
    if 'file' not in request.files:
        return jsonify({"msg": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"msg": "No selected file"}), 400
        
    if file:
        if not file.filename.lower().endswith('.pdf'):
            return jsonify({"msg": "عذراً، يُسمح فقط برفع ملفات PDF"}), 400
            
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
        if size > 5 * 1024 * 1024:
            return jsonify({"msg": "عذراً، حجم الملف يجب أن لا يتجاوز 5 ميجابايت"}), 400
            
        filename = safe_filename('doc', file.filename)
        if not filename.lower().endswith('.pdf'):
            filename = f"{filename}.pdf"
            
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        return jsonify({"msg": "تم رفع الملف بنجاح", "filename": filename}), 201

@app.route('/api/upload_photo', methods=['POST'])
@jwt_required()
def upload_photo():
    if 'file' not in request.files:
        return jsonify({"msg": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"msg": "No selected file"}), 400
        
    if file:
        if not file.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.webp')):
            return jsonify({"msg": "عذراً، يُسمح فقط برفع الصور (png, jpg, jpeg)"}), 400
            
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
        if size > 15 * 1024 * 1024:  # 15MB limit
            return jsonify({"msg": "عذراً، حجم الصورة يجب أن لا يتجاوز 15 ميجابايت"}), 400
            
        filename = secure_filename(file.filename)
        if not filename:
            import uuid
            filename = f"photo_{uuid.uuid4().hex}.jpg"
            
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        return jsonify({"msg": "تم رفع الصورة بنجاح", "photo_path": f"/api/uploads/{filename}"}), 201

@app.route('/api/documents/<filename>', methods=['GET'])
def download_document(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/api/uploads/<filename>', methods=['GET'])
def serve_upload(filename):
    import mimetypes
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.isfile(path) and open(path, 'rb').read(5) == b'%PDF-':
        return send_from_directory(app.config['UPLOAD_FOLDER'], filename, mimetype='application/pdf')
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

# --- SYSTEM LOGS ROUTE ---
@app.route('/api/logs', methods=['GET'])
@jwt_required()
@check_permission('activity_logs', 'view')
def get_logs():
    logs = SystemLog.query.order_by(SystemLog.timestamp.desc()).limit(50).all()
    return jsonify([{
        "id": l.id, 
        "timestamp": l.timestamp.isoformat(), 
        "user": l.user, 
        "action": l.action
    } for l in logs])

# --- DOCUMENT REQUESTS ROUTES ---
@app.route('/api/documents/queue', methods=['GET'])
@jwt_required()
@check_permission('doc_reupload', 'view')
def get_documents_queue():
    docs = DocumentLifecycle.query.all()
    result = []
    for d in docs:
        intern = db.session.get(Intern, d.intern_id)
        result.append({
            "id": d.id,
            "intern_id": d.intern_id,
            "intern_name": intern.name if intern else "مجهول",
            "document_type": d.doc_type,
            "custom_title": d.custom_title,
            "status": d.status,
            "file_path": d.file_path,
            "uploaded_by": d.uploaded_by,
            "uploaded_at": d.updated_at.isoformat() if d.updated_at else (d.created_at.isoformat() if d.created_at else None),
            "rejection_reason": d.rejection_reason,
            "requires_return": d.requires_return
        })
    
    return jsonify(result), 200

@app.route('/api/interns/<int:intern_id>/requests', methods=['POST'])
@jwt_required()
@check_permission('doc_reupload', 'add')
def create_document_request(intern_id):
    current_user = get_jwt()

    if request.is_json:
        data = request.json
    else:
        data = request.form

    doc_type = data.get('document_type')
    custom_title = data.get('custom_title')
    note = data.get('note')
    
    if not doc_type:
        return jsonify({"msg": "Document type is required"}), 400
        
    template_path = None
    if 'file' in request.files:
        file = request.files['file']
        if file and file.filename != '':
            if file.filename.lower().endswith('.pdf'):
                import uuid
                filename = f"tpl_{uuid.uuid4().hex}.pdf"
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                template_path = f"/api/uploads/{filename}"

    # Prevent duplicate pending requests for the same document: supersede any existing one
    DocumentRequest.query.filter_by(
        intern_id=intern_id, document_type=doc_type, custom_title=custom_title, status='pending'
    ).update({'status': 'superseded'})

    new_request = DocumentRequest(
        intern_id=intern_id,
        document_type=doc_type,
        custom_title=custom_title,
        note=note,
        status='pending',
        created_at=datetime.now(timezone.utc).isoformat(),
        template_path=template_path
    )
    db.session.add(new_request)
    db.session.commit()
    
    intern = db.session.get(Intern, intern_id)
    title_str = custom_title if doc_type == 'other' else doc_type
    log_action(current_user.get('name'), f"طلب مستند ({title_str}) من المتدرب {intern.name if intern else ''}")
    
    return jsonify({"msg": "Request created successfully", "request_id": new_request.id}), 201


@app.route('/api/interns/<int:intern_id>/document-lifecycle', methods=['POST'])
@jwt_required()
def create_intern_document_lifecycle(intern_id):
    """Admin or Intern creates a DocumentLifecycle record (request slot)."""
    current_user = get_jwt()
    role = current_user.get('role')
    if role not in ('Admin', 'Manager', 'Intern'):
        return jsonify({"msg": "Unauthorized"}), 403
    if role == 'Manager' and not has_permission(db.session.get(User, get_jwt_identity()), 'doc_reupload', 'add'):
        return jsonify({"msg": "Access Denied: لا تملك صلاحية إدارة المستندات"}), 403

    intern = db.session.get(Intern, intern_id)
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404

    data = request.json or {}
    doc_type = (data.get('document_type') or 'OTHER').strip() or 'OTHER'
    custom_title = _uniquify_title(intern_id, data.get('custom_title', '').strip() or None)
    action_type = data.get('action_type', 'view')
    now = datetime.now(timezone.utc)

    if role == 'Intern':
        requested_by = 'INTERN'
        is_intern_initiated = action_type in ('sign', 'fill', 'sign_fill')
        status = 'AWAITING_ADMIN' if is_intern_initiated else 'MISSING'
        uploaded_by = 'INTERN'
        is_visible = True
    else:
        requested_by = 'ADMIN'
        status = 'AWAITING_RETURN' if action_type in ('sign', 'fill', 'sign_fill') else 'MISSING'
        uploaded_by = 'ADMIN'
        is_visible = True

    record = DocumentLifecycle(
        intern_id=intern.id, doc_type=doc_type, status=status,
        uploaded_by=uploaded_by, is_visible_to_intern=is_visible,
        requires_return=action_type in ('sign', 'fill', 'sign_fill'),
        custom_title=custom_title, action_type=action_type,
        requested_by=requested_by,
        created_at=now, updated_at=now
    )
    db.session.add(record)
    db.session.commit()

    if role != 'Intern':
        notif_titles = {'sign': 'مستند يطلب التوقيع', 'fill': 'مستند يطلب التعبئة', 'sign_fill': 'مستند يطلب التوقيع والتعبئة', 'upload': 'مستند مطلوب رفعه', 'view': 'مستند جديد للمعاينة'}
        notif_bodies = {
            'sign': f'المستند "{custom_title or doc_type}" يتطلب منك توقيعه.',
            'fill': f'المستند "{custom_title or doc_type}" يتطلب منك تعبئته وإرجاعه.',
            'sign_fill': f'المستند "{custom_title or doc_type}" يتطلب منك توقيعه وتعبئته وإرجاعه.',
            'upload': f'مطلوب منك رفع المستند "{custom_title or doc_type}".',
            'view': f'تم رفع المستند "{custom_title or doc_type}" لمعاينته.',
        }
        notif_type = 'GENERIC' if action_type == 'view' else 'REASSIGNMENT'
        notif = Notification(intern_id=intern.id, type=notif_type,
            title=notif_titles.get(action_type, 'مستند جديد'),
            body=notif_bodies.get(action_type, 'يوجد مستند جديد.'),
            related_doc_id=record.id)
        db.session.add(notif)
        db.session.commit()

    if role == 'Intern' and action_type in ('sign', 'fill', 'sign_fill'):
        admin_notif = Notification(intern_id=intern.id, type='GENERIC',
            title=f'طلب توقيع/تعبئة من {intern.name or "متدرب"}',
            body=f'المتدرب يطلب {(custom_title or doc_type)} بتوقيع/تعبئة من الإدارة.',
            related_doc_id=record.id)
        db.session.add(admin_notif)
        db.session.commit()

    log_action(current_user.get('name') or 'متدرب', f"طلب {(custom_title or doc_type)} {'(توقيع/تعبئة)' if action_type in ('sign','fill','sign_fill') else '(عرض)'} من {'المتدرب' if role == 'Intern' else 'الإدارة'}")
    return jsonify({"msg": "تم إنشاء طلب المستند", "doc_id": record.id}), 201

@app.route('/api/intern/profile', methods=['GET'])
@jwt_required()
def get_my_intern_profile():
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"msg": "User not found"}), 404
        
    user_email = user.email
    intern = Intern.query.filter_by(email=user_email).first()
    if not intern:
        return jsonify({"msg": "Intern not found for this account"}), 404
        
    import json
    docs = {}
    if intern.documents:
        try:
            docs = json.loads(intern.documents)
        except Exception as e: print(f"Silent Error Caught: {e}")
            
    return jsonify({
        "id": intern.id, 
        "name": intern.name,
        "name_fr": intern.name_fr,
        "email": intern.email, 
        "national_id": intern.national_id,
        "department": intern.department, 
        "encadrant": intern.encadrant,
        "status": intern.status,
        "photo_path": intern.photo_path,
        "phone": intern.phone,
        "start_date": intern.start_date,
        "end_date": intern.end_date,
        "date_of_birth": intern.date_of_birth,
        "university": intern.university,
        "specialty": intern.specialty,
        "address": intern.address,
        "documents": docs
    }), 200

@app.route('/api/intern/requests', methods=['GET'])
@jwt_required()
def get_my_requests():
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"msg": "User not found"}), 404
        
    user_email = user.email
    intern = Intern.query.filter_by(email=user_email).first()
    if not intern:
        return jsonify({"msg": "Intern not found for this account"}), 404
    if intern.status == 'مرفوض':
        return jsonify({"msg": "Account not yet activated"}), 403
        
    reqs = DocumentRequest.query.filter_by(intern_id=intern.id, status='pending').all()
    
    result = []
    for r in reqs:
        result.append({
            "id": r.id,
            "document_type": r.document_type,
            "custom_title": r.custom_title,
            "note": r.note,
            "created_at": r.created_at,
            "template_path": r.template_path
        })
        
    return jsonify(result), 200

@app.route('/api/intern/requests/<int:request_id>/upload', methods=['POST'])
@jwt_required()
def upload_requested_document(request_id):
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"msg": "User not found"}), 404
        
    user_email = user.email
    
    intern = Intern.query.filter_by(email=user_email).first()
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404
    if intern.status == 'مرفوض':
        return jsonify({"msg": "Account not yet activated"}), 403
        
    doc_request = db.session.get(DocumentRequest, request_id)
    if not doc_request or doc_request.intern_id != intern.id:
        return jsonify({"msg": "Request not found"}), 404
        
    if 'file' not in request.files:
        return jsonify({"msg": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"msg": "No selected file"}), 400
        
    if file and file.filename.lower().endswith('.pdf'):
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
        if size > 15 * 1024 * 1024:
            return jsonify({"msg": "عذراً، حجم الملف يجب أن لا يتجاوز 15 ميجابايت"}), 400
            
        filename = safe_filename('req', file.filename)
        if not filename.lower().endswith('.pdf'):
            filename = f"{filename}.pdf"
            
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        file_url = f"/api/uploads/{filename}"

        try:
            import json
            try:
                docs = json.loads(intern.documents or "{}")
            except Exception:
                docs = {}
            if not isinstance(docs, dict):
                docs = {}

            if doc_request.document_type == 'other':
                if 'others' not in docs or not isinstance(docs['others'], list):
                    docs['others'] = []
                req_name = doc_request.custom_title or 'مستند إضافي'
                existing = next((o for o in docs['others'] if isinstance(o, dict) and o.get('name') == req_name), None)
                if existing:
                    existing['file'] = file_url
                else:
                    docs['others'].append({"name": req_name, "file": file_url})
            else:
                docs[doc_request.document_type] = file_url

            intern.documents = json.dumps(docs)
            
            # Fulfill this request and any other pending request for the same document
            reqs = DocumentRequest.query.filter_by(
                intern_id=intern.id,
                document_type=doc_request.document_type,
                custom_title=doc_request.custom_title,
                status='pending'
            ).all()
            for req in reqs:
                req.status = 'fulfilled'
                if req.lifecycle_id:
                    dl = db.session.get(DocumentLifecycle, req.lifecycle_id)
                    if dl:
                        old_status = dl.status
                        dl.status = 'PENDING_REVIEW'
                        dl.file_path = file_url
                        dl.uploaded_by = 'INTERN'
                        dl.rejection_reason = None
                        audit = AuditEvent(actor_user_id=user.id, actor_role='Intern',
                                           action='REUPLOAD_DOCUMENT', doc_id=dl.id, from_status=old_status,
                                           to_status='PENDING_REVIEW', note='Re-uploaded requested document')
                        db.session.add(audit)
                        
            doc_request.status = 'fulfilled'
            db.session.commit()
            log_action(user.name, f"قام برفع مستند ({doc_request.custom_title or doc_request.document_type}) استجابة لطلب")
        except Exception as e:
            db.session.rollback()
            print(f"Upload metadata update failed but file saved: {e}")
            # File was saved successfully; still report success to the client
            try:
                doc_request.status = 'fulfilled'
                db.session.commit()
            except Exception as e:
                print(f"Silent Error Caught: {e}")

        return jsonify({"msg": "تم رفع المستند بنجاح"}), 200
        
    return jsonify({"msg": "Invalid file. Only PDF is allowed."}), 400

@app.route('/api/intern/upload_unrequested', methods=['POST'])
@jwt_required()
def upload_unrequested_document():
    user_id = get_jwt_identity()
    current_user = get_jwt()
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"msg": "User not found"}), 404
        
    user_email = user.email
    
    intern = Intern.query.filter_by(email=user_email).first()
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404
    if intern.status == 'مرفوض':
        return jsonify({"msg": "Account not yet activated"}), 403
        
    doc_type = request.form.get('document_type')
    if not doc_type:
        return jsonify({"msg": "Document type is required"}), 400
        
    if 'file' not in request.files:
        return jsonify({"msg": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"msg": "No selected file"}), 400
        
    if file and file.filename.lower().endswith('.pdf'):
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
        if size > 15 * 1024 * 1024:
            return jsonify({"msg": "عذراً، حجم الملف يجب أن لا يتجاوز 15 ميجابايت"}), 400
            
        filename = safe_filename('doc', file.filename)
        if not filename.lower().endswith('.pdf'):
            filename = f"{filename}.pdf"
            
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        file_url = f"/api/uploads/{filename}"
        
        import json
        docs = json.loads(intern.documents or "{}")
        docs[doc_type] = file_url

        # Fulfill any pending request that matches this document (proactive upload counts as fulfilling)
        DocumentRequest.query.filter_by(
            intern_id=intern.id, document_type=doc_type, status='pending'
        ).update({'status': 'fulfilled'})

        intern.documents = json.dumps(docs)
        db.session.commit()

        log_action(current_user.get('name'), f"قام برفع مستند إضافي ({doc_type})")
        return jsonify({"msg": "تم رفع المستند بنجاح"}), 200
        
    return jsonify({"msg": "Invalid file. Only PDF is allowed."}), 400

# --- NOTIFICATIONS & AUDIT ---
@app.route('/api/intern/notifications', methods=['GET'])
@jwt_required()
def get_my_notifications():
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)
    if not user: return jsonify({"msg": "User not found"}), 404
    intern = Intern.query.filter_by(email=user.email).first()
    if not intern: return jsonify({"msg": "Intern not found"}), 404
    
    nots = Notification.query.filter_by(intern_id=intern.id).order_by(Notification.created_at.desc()).all()
    return jsonify([{
        "id": n.id, "type": n.type, "title": n.title, "body": n.body,
        "related_doc_id": n.related_doc_id, "is_read": n.is_read,
        "created_at": n.created_at.isoformat() if n.created_at else None
    } for n in nots])

@app.route('/api/intern/notifications/<int:nid>/read', methods=['POST'])
@jwt_required()
def read_notification(nid):
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)
    intern = Intern.query.filter_by(email=user.email).first()
    if not intern: return jsonify({"msg": "Intern not found"}), 404
    
    notif = db.session.get(Notification, nid)
    if not notif or notif.intern_id != intern.id:
        return jsonify({"msg": "Not found"}), 404
        
    notif.is_read = True
    db.session.commit()
    return jsonify({"success": True})

@app.route('/api/admin/pending-sign-fill-count', methods=['GET'])
@jwt_required()
@check_permission('statistics', 'view')
def pending_sign_fill_count():
    count = DocumentLifecycle.query.filter(
        DocumentLifecycle.action_type.in_(['sign', 'fill', 'sign_fill']),
        DocumentLifecycle.file_path.isnot(None),
        (DocumentLifecycle.returned_file_path.is_(None) | (DocumentLifecycle.returned_file_path == ''))
    ).with_entities(DocumentLifecycle.intern_id).distinct().count()
    return jsonify({"count": count})


@app.route('/api/admin/pending-review-count', methods=['GET'])
@jwt_required()
@check_permission('statistics', 'view')
def pending_review_count():
    from sqlalchemy import or_, and_
    count = DocumentLifecycle.query.filter(
        or_(
            DocumentLifecycle.status == 'AWAITING_ADMIN',
            and_(
                DocumentLifecycle.status.in_(['PENDING_REVIEW', 'RETURNED']),
                or_(
                    DocumentLifecycle.uploaded_by != 'ADMIN',
                    DocumentLifecycle.uploaded_by == None
                )
            ),
            and_(
                DocumentLifecycle.status == 'REVISION_REQUESTED',
                DocumentLifecycle.revision_requested_by == 'INTERN'
            )
        )
    ).with_entities(DocumentLifecycle.intern_id).distinct().count()
    return jsonify({"count": count})


@app.route('/api/notifications', methods=['GET'])
@jwt_required()
def list_notifications():
    current_user = get_jwt()
    if current_user.get('role') == 'Intern':
        user = db.session.get(User, get_jwt_identity())
        intern = Intern.query.filter_by(email=user.email).first() if user else None
        if not intern:
            return jsonify([])
        notifs = Notification.query.filter_by(intern_id=intern.id).order_by(Notification.created_at.desc()).limit(50).all()
    elif current_user.get('role') in ('Admin', 'Manager'):
        notifs = Notification.query.filter(Notification.type.in_(['GENERIC', 'DOCUMENT_RETURNED'])).order_by(Notification.created_at.desc()).limit(50).all()
    else:
        return jsonify({"msg": "Unauthorized"}), 403
    return jsonify([{
        "id": n.id, "intern_id": n.intern_id, "type": n.type,
        "title": n.title, "body": n.body, "related_doc_id": n.related_doc_id,
        "is_read": n.is_read, "created_at": n.created_at
    } for n in notifs])

@app.route('/api/notifications/<int:nid>/read', methods=['POST'])
@jwt_required()
def mark_notification_read(nid):
    current_user = get_jwt()
    if current_user.get('role') not in ('Admin', 'Manager'):
        return jsonify({"msg": "Unauthorized"}), 403
    n = db.session.get(Notification, nid)
    if not n: return jsonify({"msg": "Not found"}), 404
    n.is_read = True
    db.session.commit()
    return jsonify({"success": True})


@app.route('/api/interns/<int:intern_id>/audit', methods=['GET'])
@jwt_required()
@check_permission('doc_reupload', 'view')
def get_intern_audit(intern_id):
    events = AuditEvent.query.join(DocumentLifecycle, DocumentLifecycle.id == AuditEvent.doc_id)\
            .filter(DocumentLifecycle.intern_id == intern_id).order_by(AuditEvent.created_at.desc()).all()
    return jsonify([{
        "id": e.id, "actor_role": e.actor_role, "action": e.action, "doc_id": e.doc_id,
        "from_status": e.from_status, "to_status": e.to_status, "note": e.note,
        "created_at": e.created_at.isoformat() if e.created_at else None
    } for e in events])


# --- NEW DOCUMENT WORKFLOW ROUTES ---
@app.route('/api/interns/<int:intern_id>/documents/<int:doc_id>/reject', methods=['POST'])
@jwt_required()
@check_permission('doc_reupload', 'approve')
def reject_document(intern_id, doc_id):
    current_user = get_jwt()
    data = request.json or {}
    reason = data.get('reason') or data.get('rejection_reason')
    if not reason: return jsonify({"msg": "Reason required"}), 400
    
    doc = db.session.get(DocumentLifecycle, doc_id)
    if not doc or doc.intern_id != intern_id: return jsonify({"msg": "Not found"}), 404
    
    old_status = doc.status
    doc.status = 'REVISION_REQUESTED'
    doc.correction_round += 1
    doc.rejection_reason = reason
    doc.revision_requested_by = 'ADMIN'
    
    audit = AuditEvent(actor_user_id=current_user.get('sub'), actor_role=current_user.get('role'),
                       action='REJECT_DOCUMENT', doc_id=doc.id, from_status=old_status,
                       to_status='REVISION_REQUESTED', note=reason)
    db.session.add(audit)
    
    notif = Notification(intern_id=intern_id, type='REJECTION', title='تم رفض المستند',
                         body=f'تم رفض مستند {doc.custom_title or doc.doc_type} بسبب: {reason}',
                         related_doc_id=doc.id)
    db.session.add(notif)
    Notification.query.filter_by(related_doc_id=doc.id, type='GENERIC').update({'is_read': True})
    db.session.commit()
    return jsonify({"success": True})

@app.route('/api/interns/<int:intern_id>/documents/<int:doc_id>', methods=['DELETE'])
@jwt_required()
@check_permission('doc_reupload', 'delete')
def delete_intern_document(intern_id, doc_id):
    current_user = get_jwt()
    doc = db.session.get(DocumentLifecycle, doc_id)
    if not doc or doc.intern_id != intern_id:
        return jsonify({"msg": "Document not found"}), 404

    intern = db.session.get(Intern, intern_id)
    Notification.query.filter_by(related_doc_id=doc.id).delete()
    AuditEvent.query.filter_by(doc_id=doc.id).delete()
    DocumentRequest.query.filter_by(lifecycle_id=doc.id).delete()
    DocumentLifecycle.query.filter_by(parent_id=doc.id).delete()
    db.session.delete(doc)
    db.session.commit()
    log_action(current_user.get('name'), f"حذف مستند ({doc.custom_title or doc.doc_type}) للمتدرب {intern.name if intern else ''}")
    return jsonify({"success": True}), 200

@app.route('/api/interns/<int:intern_id>/documents/<int:doc_id>/assign', methods=['POST'])
@jwt_required()
@check_permission('doc_reupload', 'edit')
def assign_document(intern_id, doc_id):
    current_user = get_jwt()
    data = request.json or {}
    assign_to_id = data.get('assign_to_id')
    if not assign_to_id: return jsonify({"msg": "assign_to_id required"}), 400
    
    doc = db.session.get(DocumentLifecycle, doc_id)
    if not doc or doc.intern_id != intern_id: return jsonify({"msg": "Not found"}), 404
    
    doc.assigned_to_intern_id = assign_to_id
    
    audit = AuditEvent(actor_user_id=current_user.get('sub'), actor_role=current_user.get('role'),
                       action='ASSIGN_DOCUMENT', doc_id=doc.id, note=f'Assigned to {assign_to_id}')
    db.session.add(audit)
    
    notif = Notification(intern_id=assign_to_id, type='REASSIGNMENT', title='تم تعيين مستند لك',
                         body=f'تم تعيين مستند {doc.custom_title or doc.doc_type} لمراجعته أو توقيعه.',
                         related_doc_id=doc.id)
    db.session.add(notif)
    db.session.commit()
    return jsonify({"success": True})

@app.route('/api/interns/<int:intern_id>/complete-stage', methods=['POST'])
@jwt_required()
@check_permission('attestation', 'approve')
def complete_stage(intern_id):
    intern = db.session.get(Intern, intern_id)
    if not intern: return jsonify({"msg": "Not found"}), 404
    
    required_types = ['CIN', 'CV', 'DEMANDE', 'CONVENTION_SIGNED', 'FINAL_REPORT']
    docs = DocumentLifecycle.query.filter_by(intern_id=intern_id).all()
    for req in required_types:
        d = next((x for x in docs if x.doc_type == req), None)
        if not d or d.status != 'APPROVED_AND_SIGNED':
            return jsonify({"msg": f"المستند {req} غير مكتمل."}), 400
            
    notif = Notification(intern_id=intern_id, type='STAGE_COMPLETE', title='تهانينا!',
                         body='تم إنهاء التدريب بنجاح وإصدار الشهادة.')
    db.session.add(notif)
    
    intern.status = 'مكتمل'
    db.session.commit()
    
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        import io
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        c.setFont("Helvetica-Bold", 24)
        c.drawCentredString(width/2.0, height - 100, "ATTESTATION DE STAGE")
        c.setFont("Helvetica", 14)
        c.drawString(50, height - 200, f"Nous soussignes, certifions que Monsieur/Madame:")
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 230, f"{intern.name_fr or '__________________'}")
        c.setFont("Helvetica", 14)
        c.drawString(50, height - 280, f"A effectue un stage au sein de notre departement:")
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, height - 310, f"{intern.department or '__________________'}")
        c.setFont("Helvetica", 14)
        c.drawString(50, height - 360, f"Encadre(e) par:")
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, height - 390, f"{intern.encadrant or '__________________'}")
        c.setFont("Helvetica", 14)
        c.drawString(50, height - 440, f"Période: du {intern.start_date or '___'} au {intern.end_date or '___'}")
        c.drawString(50, height - 500, "Cette attestation est delivree pour servir et valoir ce que de droit.")
        c.drawString(width - 200, 150, "Signature et Cachet:")
        c.showPage()
        c.save()
        
        buffer.seek(0)
        filename = f"Attestation_Signed_{intern.id}.pdf"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        with open(filepath, 'wb') as f:
            f.write(buffer.read())
            
        file_url = f"/api/uploads/{filename}"
        dl = DocumentLifecycle(
            intern_id=intern.id, doc_type='ATTESTATION_SIGNED',
            file_path=file_url, uploaded_by='ADMIN',
            status='APPROVED_AND_SIGNED', is_visible_to_intern=True,
            custom_title='شهادة التدريب', created_at=datetime.utcnow()
        )
        db.session.add(dl)
        db.session.commit()
    except Exception as e:
        print(f"Failed to generate certificate: {e}")
        
    return jsonify({"success": True})

# --- MESSAGING ROUTES (admin <-> intern) ---
def build_file_url(path: str) -> str:
    if not path:
        return ''
    if path.startswith('http'):
        return path
    name = path.replace('/api/uploads/', '').replace('/api/documents/', '').replace('/', '')
    return f"http://127.0.0.1:5055/api/uploads/{name}"

def allowed_message_attachment(filename: str) -> bool:
    if not filename:
        return False
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    return ext in ('pdf', 'doc', 'docx', 'txt', 'png', 'jpg', 'jpeg', 'webp')

@app.route('/api/interns/<int:intern_id>/messages', methods=['GET'])
@jwt_required()
def get_messages(intern_id):
    msgs = Message.query.filter_by(intern_id=intern_id).order_by(Message.created_at.asc()).all()
    result = []
    for m in msgs:
        result.append({
            "id": m.id,
            "intern_id": m.intern_id,
            "sender_role": m.sender_role,
            "sender_name": m.sender_name,
            "body": m.body,
            "attachment_path": m.attachment_path,
            "attachment_name": m.attachment_name,
            "attachment_url": build_file_url(m.attachment_path) if m.attachment_path else None,
            "waiting_for_reply": m.waiting_for_reply,
            "expected_format": m.expected_format,
            "replied": m.replied,
            "created_at": m.created_at
        })
    return jsonify(result), 200

@app.route('/api/interns/<int:intern_id>/messages', methods=['POST'])
@jwt_required()
def send_message(intern_id):
    current_user = get_jwt()
    intern = db.session.get(Intern, intern_id)
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404

    role = current_user.get('role')
    if role == 'Intern':
        user_id = get_jwt_identity()
        user = db.session.get(User, user_id)
        if not user or (intern.email and user.email != intern.email):
            return jsonify({"msg": "Unauthorized"}), 403

    is_json = request.is_json
    body = (request.form.get('body') if not is_json else request.json.get('body')) or ''
    waiting_for_reply = False
    expected_format = None

    if role != 'Intern':
        waiting_for_reply = str((request.form.get('waiting_for_reply') if not is_json else request.json.get('waiting_for_reply')) or '').lower() in ('1', 'true', 'on')
        expected_format = (request.form.get('expected_format') if not is_json else request.json.get('expected_format')) or None

    attachment_path = None
    attachment_name = None
    if 'file' in request.files:
        file = request.files['file']
        if file and file.filename != '':
            if not allowed_message_attachment(file.filename):
                return jsonify({"msg": "عذراً، صيغة الملف غير مدعومة"}), 400
            file.seek(0, os.SEEK_END)
            size = file.tell()
            file.seek(0)
            if size > 15 * 1024 * 1024:
                return jsonify({"msg": "عذراً، حجم الملف يجب أن لا يتجاوز 15 ميجابايت"}), 400
            filename = safe_filename('msg', file.filename)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            attachment_path = f"/api/uploads/{filename}"
            attachment_name = file.filename

    # When an intern replies to a message, mark waiting messages as replied
    if role == 'Intern' and (body.strip() != '' or attachment_path):
        Message.query.filter_by(intern_id=intern_id, waiting_for_reply=True).update({'replied': True})

    new_msg = Message(
        intern_id=intern_id,
        sender_role='intern' if role == 'Intern' else 'admin',
        sender_name=current_user.get('name'),
        body=body,
        attachment_path=attachment_path,
        attachment_name=attachment_name,
        waiting_for_reply=waiting_for_reply and role != 'Intern',
        expected_format=expected_format if role != 'Intern' else None,
        replied=False,
        created_at=datetime.now(timezone.utc).isoformat()
    )
    db.session.add(new_msg)
    db.session.commit()

    log_action(current_user.get('name'), f"أرسل رسالة للمتدرب {intern.name}")
    return jsonify({
        "id": new_msg.id,
        "intern_id": new_msg.intern_id,
        "sender_role": new_msg.sender_role,
        "sender_name": new_msg.sender_name,
        "body": new_msg.body,
        "attachment_path": new_msg.attachment_path,
        "attachment_name": new_msg.attachment_name,
        "attachment_url": build_file_url(new_msg.attachment_path) if new_msg.attachment_path else None,
        "waiting_for_reply": new_msg.waiting_for_reply,
        "expected_format": new_msg.expected_format,
        "replied": new_msg.replied,
        "created_at": new_msg.created_at
    }), 201

@app.route('/api/interns/<int:intern_id>/messages/<int:message_id>', methods=['DELETE'])
@jwt_required()
def delete_message(intern_id, message_id):
    current_user = get_jwt()
    if current_user.get('role') != 'Admin':
        return jsonify({"msg": "Unauthorized"}), 403
    msg = Message.query.filter_by(id=message_id, intern_id=intern_id).first()
    if not msg:
        return jsonify({"msg": "Message not found"}), 404
    if msg.attachment_path:
        try:
            name = msg.attachment_path.replace('/api/uploads/', '').replace('/', '')
            fp = os.path.join(app.config['UPLOAD_FOLDER'], name)
            if os.path.exists(fp):
                os.remove(fp)
        except Exception as e:
                print(f"Silent Error Caught: {e}")
    db.session.delete(msg)
    db.session.commit()
    return jsonify({"success": True}), 200

# --- PERSONAL PROFILE EXPORT (md / PDF / Excel) ---
def render_intern_md(intern: Intern) -> str:
    docs = {}
    if intern.documents:
        try:
            docs = json.loads(intern.documents)
        except Exception:
            docs = {}
    others = docs.get('others', []) if isinstance(docs.get('others'), list) else []

    def row(label, val):
        return f"| {label} | {val or '—'} |"

    lines = []
    lines.append(f"# ملف المتدرب: {intern.name}")
    if intern.name_fr:
        lines.append(f"**{intern.name_fr}**")
    lines.append("")
    lines.append("## المعلومات الشخصية")
    lines.append("")
    lines.append("| الحقل | القيمة |")
    lines.append("| --- | --- |")
    lines.append(row("رقم التسجيل", f"INT-{intern.id:04d}"))
    lines.append(row("رقم الهوية الوطنية", intern.national_id))
    lines.append(row("البريد الإلكتروني", intern.email))
    lines.append(row("رقم الهاتف", intern.phone))
    lines.append(row("تاريخ الازدياد", intern.date_of_birth))
    lines.append(row("تاريخ البدء", intern.start_date))
    lines.append(row("تاريخ الانتهاء", intern.end_date))
    lines.append(row("الجامعة أو المعهد", intern.university))
    lines.append(row("التخصص", intern.specialty))
    lines.append(row("القسم", intern.department))
    lines.append(row("المؤطر", intern.encadrant))
    lines.append(row("الحالة", intern.status))
    lines.append(row("العنوان", intern.address))
    lines.append("")
    lines.append("## المستندات")
    lines.append("")
    if docs:
        for k, v in docs.items():
            if k == 'others':
                continue
            if v:
                lines.append(f"- {k}: {v}")
    for o in others:
        if isinstance(o, dict) and o.get('file'):
            lines.append(f"- {o.get('name', 'مستند إضافي')}: {o.get('file')}")
    if not docs and not others:
        lines.append("- لا توجد مستندات.")
    lines.append("")
    return "\n".join(lines)

@app.route('/api/interns/<int:intern_id>/profile-md', methods=['GET'])
@jwt_required()
@check_permission('interns', 'approve')
def download_profile_md(intern_id):
    intern = db.session.get(Intern, intern_id)
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404
    md = render_intern_md(intern)
    buffer = io.BytesIO(md.encode('utf-8'))
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f"Profil_{intern.id}.md", mimetype='text/markdown')

@app.route('/api/interns/<int:intern_id>/profile-pdf', methods=['GET'])
@jwt_required()
@check_permission('interns', 'approve')
def download_profile_pdf(intern_id):
    intern = db.session.get(Intern, intern_id)
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404
    mode = request.args.get('mode', 'summary').lower()
    if mode not in ('summary', 'full'):
        mode = 'summary'
    disposition = request.args.get('disposition', 'attachment').lower()
    as_attachment = disposition != 'inline'
    try:
        from pdf_report import build_intern_pdf, build_filename
        buffer = build_intern_pdf([intern], mode=mode)
        filename = build_filename(intern)
    except ImportError:
        return jsonify({"msg": "PDF generation library not installed"}), 500
    return send_file(buffer, as_attachment=as_attachment, download_name=filename, mimetype='application/pdf')

@app.route('/api/interns/export', methods=['GET'])
@jwt_required()
@check_permission('interns', 'approve')
def export_interns():
    fmt = request.args.get('format', 'pdf').lower()
    ids_param = request.args.get('ids')
    if ids_param:
        try:
            ids = [int(x) for x in ids_param.split(',') if x.strip()]
        except ValueError:
            return jsonify({"msg": "Invalid ids"}), 400
        interns = Intern.query.filter(Intern.id.in_(ids)).all()
    else:
        interns = Intern.query.all()

    if not interns:
        return jsonify({"msg": "No interns to export"}), 404

    if fmt == 'excel':
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment
            from openpyxl.utils import get_column_letter
        except ImportError:
            return jsonify({"msg": "Excel library not installed"}), 500

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Interns"
        headers = ["ID", "الاسم", "الاسم (فرنسية)", "البريد", "الهاتف", "رقم الهوية",
                   "القسم", "المؤطر", "الحالة", "تاريخ البدء", "تاريخ الانتهاء",
                   "تاريخ الازدياد", "الجامعة", "التخصص", "العنوان"]
        ws.append(headers)
        header_fill = PatternFill(start_color="1E5631", end_color="1E5631", fill_type="solid")
        for col, _ in enumerate(headers, start=1):
            c = ws.cell(row=1, column=col)
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = header_fill
            c.alignment = Alignment(horizontal="right", vertical="center")
        ws.sheet_view.rightToLeft = True
        for i in interns:
            row_data = [
                i.id, i.name, i.name_fr, i.email, i.phone, i.national_id,
                i.department, i.encadrant, i.status, i.start_date, i.end_date,
                i.date_of_birth, i.university, i.specialty, i.address
            ]
            ws.append(row_data)
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=len(headers)):
            for cell in row:
                cell.alignment = Alignment(horizontal="right", vertical="center")
        for col in range(1, len(headers) + 1):
            ws.column_dimensions[get_column_letter(col)].width = 18
        ws.column_dimensions['B'].width = 24
        ws.column_dimensions['N'].width = 30

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name="Interns_Export.xlsx", mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

    # Default: PDF (one formal profile per page)
    try:
        from pdf_report import build_intern_pdf, build_filename
        buffer = build_intern_pdf(interns)
        filename = build_filename(interns[0]) if len(interns) == 1 else "Interns_Export.pdf"
    except ImportError:
        return jsonify({"msg": "PDF generation library not installed"}), 500
    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/pdf')


# --- DOCUMENT LIFECYCLE ENDPOINTS ---


def _get_doc_type_intern():
    """For intern-portal endpoints: return (intern, user_claims) from JWT matching Intern.email."""
    claims = get_jwt()
    role = claims.get('role')
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)
    if not user:
        return None, None, None
    if role == 'Intern':
        intern = Intern.query.filter_by(email=user.email).first()
        return intern, claims, user
    return None, claims, user


@app.route('/api/interns/<int:intern_id>/documents', methods=['GET'])
@jwt_required()
def list_intern_documents(intern_id):
    intern = db.session.get(Intern, intern_id)
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404
    docs = DocumentLifecycle.query.filter_by(intern_id=intern_id).order_by(DocumentLifecycle.created_at.asc()).all()
    result = []
    for d in docs:
        result.append({
            "id": d.id,
            "doc_type": d.doc_type,
            "label": d.custom_title or d.doc_type,
            "file_path": d.file_path,
            "uploaded_by": d.uploaded_by,
            "status": d.status,
            "rejection_reason": d.rejection_reason,
            "is_visible_to_intern": d.is_visible_to_intern,
            "custom_title": d.custom_title,
            "requires_return": d.requires_return,
            "returned_file_path": d.returned_file_path,
            "returned_files_history": d.returned_files_history,
            "created_at": d.created_at,
            "updated_at": d.updated_at,
            "file_type": d.file_type or 'pdf',
            "action_type": d.action_type or 'view',
            "requested_by": d.requested_by or ('INTERN' if (d.uploaded_by or '') == 'INTERN' else 'ADMIN'),
            "revision_requested_by": d.revision_requested_by,
            "source": d.source or 'ADMIN',
        })
    return jsonify(result), 200


# ──────────────────────────────────────────────
# Document Templates (admin-managed required docs)
# ──────────────────────────────────────────────
@app.route('/api/admin/document-templates', methods=['GET'])
@jwt_required()
@check_permission('doc_templates', 'view')
def get_document_templates():
    templates = DocumentTemplate.query.order_by(DocumentTemplate.id).all()
    return jsonify([{
        "id": t.id, "label": t.label,
        "file_type": t.file_type, "file_path": t.file_path,
        "is_active": t.is_active,
        "created_at": t.created_at
    } for t in templates])

@app.route('/api/admin/document-templates', methods=['POST'])
@jwt_required()
@check_permission('doc_templates', 'add')
def create_document_template():
    json_data = request.get_json(silent=True) or {}
    label = (request.form.get('label') or json_data.get('label') or '').strip()
    if not label:
        return jsonify({"msg": "Label is required"}), 400
    file_type = request.form.get('file_type') or json_data.get('file_type') or 'pdf'
    file_path = None
    if 'file' in request.files:
        file = request.files['file']
        if file.filename:
            import re, random, string
            ext = ('.' + file.filename.rsplit('.', 1)[1].lower()) if '.' in file.filename else '.pdf'
            ts = datetime.now().strftime('%Y%m%d%H%M%S')
            rand = ''.join(random.choices(string.ascii_lowercase + string.digits, k=8))
            safe = re.sub(r'[^\w\s-]', '', label).strip().replace(' ', '-')[:30].lower() or 'doc'
            filename = f"template-{ts}-{rand}-{safe}{ext}"
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            file_path = f"/api/uploads/{filename}"
    t = DocumentTemplate(label=label, file_type=file_type, file_path=file_path, is_active=True)
    db.session.add(t)
    db.session.commit()
    return jsonify({"id": t.id, "file_path": file_path, "msg": "Template created"}), 201

@app.route('/api/admin/document-templates/<int:tid>', methods=['PUT'])
@jwt_required()
@check_permission('doc_templates', 'edit')
def update_document_template(tid):
    t = db.session.get(DocumentTemplate, tid)
    if not t: return jsonify({"msg": "Not found"}), 404
    json_data = request.get_json(silent=True) or {}
    label = request.form.get('label') or json_data.get('label')
    if label: t.label = label.strip()
    
    # Check keys safely
    has_file_type = 'file_type' in request.form or 'file_type' in json_data
    if has_file_type:
        t.file_type = request.form.get('file_type') or json_data.get('file_type') or t.file_type
        
    has_is_active = 'is_active' in request.form or 'is_active' in json_data
    if has_is_active:
        if 'is_active' in request.form:
            t.is_active = request.form.get('is_active', type=int)
        else:
            t.is_active = json_data.get('is_active', t.is_active)
            
    if 'file' in request.files:
        file = request.files['file']
        if file.filename:
            import re, random, string
            ext = ('.' + file.filename.rsplit('.', 1)[1].lower()) if '.' in file.filename else '.pdf'
            ts = datetime.now().strftime('%Y%m%d%H%M%S')
            rand = ''.join(random.choices(string.ascii_lowercase + string.digits, k=8))
            safe = re.sub(r'[^\w\s-]', '', t.label).strip().replace(' ', '-')[:30].lower() or 'doc'
            filename = f"template-{ts}-{rand}-{safe}{ext}"
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            t.file_path = f"/api/uploads/{filename}"
    db.session.commit()
    return jsonify({"msg": "Updated", "file_path": t.file_path})

@app.route('/api/admin/document-templates/<int:tid>', methods=['DELETE'])
@jwt_required()
@check_permission('doc_templates', 'delete')
def delete_document_template(tid):
    t = db.session.get(DocumentTemplate, tid)
    if not t: return jsonify({"msg": "Not found"}), 404
    db.session.delete(t)
    db.session.commit()
    return jsonify({"msg": "Deleted"})
# ──────────────────────────────────────────────


@app.route('/api/interns/<int:intern_id>/documents/upload', methods=['POST'])
@jwt_required()
def upload_intern_document(intern_id):
    """Intern uploads a file for a given doc_type (creates or updates lifecycle record)."""
    intern, claims, user = _get_doc_type_intern()
    if not intern or intern.id != intern_id:
        current_user = get_jwt()
        if current_user.get('role') not in ('Admin',):
            return jsonify({"msg": "Unauthorized"}), 403
        intern = db.session.get(Intern, intern_id)
        if not intern:
            return jsonify({"msg": "Intern not found"}), 404
    elif intern.status == 'مرفوض' and claims.get('role') == 'Intern':
        return jsonify({"msg": "Account not yet activated"}), 403

    doc_type = (request.form.get('doc_type') or 'OTHER').strip() or 'OTHER'
    # DOC_TYPES is now only a legacy reference set; arbitrary types are allowed.
    if not doc_type or len(doc_type) > 50:
        return jsonify({"msg": "Invalid doc_type"}), 400
    target_doc_id = request.form.get('doc_id')

    if 'file' not in request.files:
        return jsonify({"msg": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"msg": "No selected file"}), 400

    custom_title = request.form.get('custom_title')
    now = datetime.now(timezone.utc)

    if target_doc_id:
        record = DocumentLifecycle.query.filter_by(
            id=target_doc_id, intern_id=intern.id
        ).first()
        if not record:
            return jsonify({"msg": "Document slot not found"}), 404
    else:
        existing = DocumentLifecycle.query.filter_by(
            intern_id=intern.id, doc_type=doc_type
        ).filter(
            DocumentLifecycle.custom_title.is_(None) if not custom_title else DocumentLifecycle.custom_title == custom_title
        ).first()
        record = existing

    file_type = (record.file_type if record else request.form.get('file_type')) or 'pdf'
    ext = ('.' + file.filename.rsplit('.', 1)[1].lower()) if '.' in file.filename else ''
    if not record and not request.form.get('file_type'):
        for ft, exts in FILE_TYPE_EXTENSIONS.items():
            if ext in exts and ft != 'any':
                file_type = ft
                break
    elif not _allowed_file(file.filename, file_type):
        for ft, exts in FILE_TYPE_EXTENSIONS.items():
            if ext in exts:
                file_type = ft
                break
    if not _allowed_file(file.filename, file_type):
        return jsonify({"msg": f"File type not allowed (expected: {file_type})"}), 400
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > 15 * 1024 * 1024:
        return jsonify({"msg": "File too large (max 15MB)"}), 400

    label_for_name = (record.custom_title if record else custom_title) or doc_type

    if not record:
        record = DocumentLifecycle(
            intern_id=intern.id, doc_type=doc_type,
            status='APPROVED_AND_SIGNED', # Since proactive upload defaults to 'view'
            action_type='view',
            requested_by='INTERN',
            uploaded_by='INTERN', is_visible_to_intern=True,
            custom_title=custom_title, file_type=file_type,
            created_at=now, updated_at=now
        )
        db.session.add(record)
        db.session.flush()
    else:
        # Preserve AWAITING_ADMIN status for intern-initiated sign/fill requests
        if record.action_type == 'view':
            record.status = 'APPROVED_AND_SIGNED'
            record.is_visible_to_intern = True
        elif record.requested_by == 'INTERN' and record.action_type in ('sign', 'fill', 'sign_fill') and record.status in ('AWAITING_ADMIN', 'REVISION_REQUESTED'):
            record.status = 'AWAITING_ADMIN'
        else:
            record.status = 'PENDING_REVIEW'
        record.uploaded_by = 'INTERN'
        record.updated_at = now

    safe_label = re.sub(r'[^\w\s-]', '', label_for_name).strip().replace(' ', '-')[:40].lower() or 'doc'
    orig_ext = ('.' + file.filename.rsplit('.', 1)[1].lower()) if '.' in file.filename else ''
    if not orig_ext or len(orig_ext) > 10 or not orig_ext[1:].isalnum():
        orig_ext = '.pdf'
    suffix = f"intern-{intern.id}-{record.id}-{safe_label}{orig_ext}"
    file.save(os.path.join(app.config['UPLOAD_FOLDER'], suffix))
    file_url = f"/api/uploads/{suffix}"

    record.file_path = file_url
    if not record.id:
        db.session.flush()
    else:
        record.updated_at = now
        record.rejection_reason = None
        if file_type:
            record.file_type = file_type
    audit = AuditEvent(actor_user_id=user.id if user else None, actor_role=claims.get('role') if claims else 'Intern',
                       action='UPLOAD_DOCUMENT', doc_id=record.id, from_status=record.status,
                       to_status='PENDING_REVIEW', note='Intern re-uploaded document')
    db.session.add(audit)
    db.session.commit()

    # Notify admins about the upload
    doc_title = (record.custom_title if record else custom_title) or doc_type
    admin_notif = Notification(intern_id=intern.id, type='GENERIC',
        title=f'مستند برفع من {intern.name or "متدرب"}',
        body=f'تم رفع "{doc_title}" من قبل المتدرب وهو قيد المراجعة.',
        related_doc_id=record.id)
    db.session.add(admin_notif)
    db.session.commit()

    # Auto-fulfill any pending DocumentRequest for this doc_type
    if not custom_title:
        DocumentRequest.query.filter_by(
            intern_id=intern.id, document_type=doc_type, status='pending'
        ).update({'status': 'fulfilled'})

    return jsonify({"msg": "Uploaded successfully", "doc": {
        "id": record.id, "doc_type": record.doc_type, "status": record.status,
        "file_path": record.file_path
    }}), 200


@app.route('/api/interns/<int:intern_id>/documents/<int:doc_id>/approve', methods=['POST'])
@jwt_required()
@check_permission('doc_reupload', 'approve')
def approve_document(intern_id, doc_id):
    current_user = get_jwt()
    doc = DocumentLifecycle.query.filter_by(id=doc_id, intern_id=intern_id).first()
    if not doc:
        return jsonify({"msg": "Document not found"}), 404
    doc.status = 'APPROVED_AND_SIGNED'
    doc.is_visible_to_intern = True
    doc.updated_at = datetime.now(timezone.utc)
    Notification.query.filter_by(related_doc_id=doc.id).update({'is_read': True})
    db.session.commit()
    intern = db.session.get(Intern, intern_id)
    log_action(current_user.get('name'), f"قبول مستند ({doc.doc_type}) للمتدرب {intern.name if intern else ''}")
    return jsonify({"success": True, "status": doc.status}), 200




@app.route('/api/interns/<int:intern_id>/documents/signed', methods=['POST'])
@jwt_required()
@check_permission('doc_reupload', 'add')
def upload_signed_document(intern_id):
    """Admin uploads a signed/approved version of a document (outgoing to intern)."""
    current_user = get_jwt()
    intern = db.session.get(Intern, intern_id)
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404

    doc_type = (request.form.get('doc_type') or 'OTHER').strip() or 'OTHER'
    if len(doc_type) > 50:
        return jsonify({"msg": "Invalid doc_type"}), 400

    custom_title = request.form.get('custom_title', '').strip() or None
    action_type = (request.form.get('action_type') or '').strip() or None
    target_doc_id = request.form.get('doc_id')
    file_type = (request.form.get('file_type') or '').strip().lower() or None

    if 'file' not in request.files:
        return jsonify({"msg": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"msg": "No selected file"}), 400

    if target_doc_id:
        existing = DocumentLifecycle.query.filter_by(id=target_doc_id, intern_id=intern.id).first()
        if not existing:
            return jsonify({"msg": "Document slot not found"}), 404
    elif custom_title:
        existing = DocumentLifecycle.query.filter_by(
            intern_id=intern.id, doc_type=doc_type, custom_title=custom_title, action_type=action_type or 'view'
        ).first()
    else:
        existing = DocumentLifecycle.query.filter_by(intern_id=intern.id, doc_type=doc_type, action_type=action_type or 'view').filter(
            DocumentLifecycle.custom_title.is_(None)
        ).first()

    record_ft = (existing.file_type if existing else file_type) or 'pdf'
    if not file_type and not (existing and existing.file_type):
        ext = ('.' + file.filename.rsplit('.', 1)[1].lower()) if '.' in file.filename else ''
        for ft, exts in FILE_TYPE_EXTENSIONS.items():
            if ext in exts and ft != 'any':
                record_ft = ft
                break
    if not _allowed_file(file.filename, record_ft):
        return jsonify({"msg": f"File type not allowed (expected: {record_ft})"}), 400
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > 15 * 1024 * 1024:
        return jsonify({"msg": "File too large (max 15MB)"}), 400

    filename = safe_filename('signed', file.filename)
    file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
    file_url = f"/api/uploads/{filename}"

    now = datetime.now(timezone.utc)

    if existing:
        existing.file_path = file_url
        existing.status = 'AWAITING_RETURN' if (existing.action_type in ('sign', 'fill', 'sign_fill')) else 'APPROVED_AND_SIGNED'
        existing.uploaded_by = 'ADMIN'
        existing.is_visible_to_intern = True
        if custom_title:
            existing.custom_title = custom_title
        if file_type:
            existing.file_type = file_type
        existing.updated_at = now
        db.session.commit()
        
        notif = Notification(intern_id=intern.id, type='GENERIC',
            title='تحديث على مستند',
            body=f'تم تحديث المستند "{custom_title or doc_type}" من قبل الإدارة.',
            related_doc_id=existing.id)
        db.session.add(notif)
        db.session.commit()
        
        log_action(current_user.get('name'), f"رفع نسخة موقعة من {doc_type} للمتدرب {intern.name}")
        return jsonify({"msg": "Signed document uploaded", "doc_id": existing.id}), 200

    record = DocumentLifecycle(
        intern_id=intern.id, doc_type=doc_type, file_path=file_url,
        uploaded_by='ADMIN', status='AWAITING_RETURN' if (action_type in ('sign', 'fill', 'sign_fill')) else 'APPROVED_AND_SIGNED',
        is_visible_to_intern=True, custom_title=_uniquify_title(intern.id, custom_title),
        file_type=file_type or 'pdf',
        action_type=action_type or 'view',
        created_at=now, updated_at=now
    )
    db.session.add(record)
    db.session.commit()
    
    notif_titles = {'sign': 'مستند يطلب التوقيع', 'fill': 'مستند يطلب التعبئة', 'sign_fill': 'مستند يطلب التوقيع والتعبئة', 'upload': 'مستند مطلوب رفعه', 'view': 'مستند جديد للمعاينة'}
    notif_bodies = {
        'sign': f'المستند "{custom_title or doc_type}" يتطلب منك توقيعه.',
        'fill': f'المستند "{custom_title or doc_type}" يتطلب منك تعبئته وإرجاعه.',
        'sign_fill': f'المستند "{custom_title or doc_type}" يتطلب منك توقيعه وتعبئته وإرجاعه.',
        'upload': f'مطلوب منك رفع المستند "{custom_title or doc_type}".',
        'view': f'تم إرفاق المستند "{custom_title or doc_type}" لمعاينته.',
    }
    notif_type = 'GENERIC' if action_type == 'view' else 'REASSIGNMENT'
    notif = Notification(intern_id=intern.id, type=notif_type,
        title=notif_titles.get(action_type or 'view', 'مستند جديد'),
        body=notif_bodies.get(action_type or 'view', 'يوجد مستند جديد.'),
        related_doc_id=record.id)
    db.session.add(notif)
    db.session.commit()
    
    log_action(current_user.get('name'), f"رفع نسخة موقعة من {doc_type} للمتدرب {intern.name}")
    return jsonify({"msg": "Signed document uploaded", "doc_id": record.id}), 201


@app.route('/api/interns/<int:intern_id>/documents/<int:doc_id>/return-upload', methods=['POST'])
@jwt_required()
def upload_returned_document(intern_id, doc_id):
    """Intern uploads the filled version of a requires_return document."""
    intern, claims, user = _get_doc_type_intern()
    if not intern or intern.id != intern_id:
        current_user = get_jwt()
        if current_user.get('role') not in ('Admin',):
            return jsonify({"msg": "Unauthorized"}), 403
        intern = db.session.get(Intern, intern_id)
        if not intern:
            return jsonify({"msg": "Intern not found"}), 404

    doc = DocumentLifecycle.query.filter_by(id=doc_id, intern_id=intern_id).first()
    if not doc:
        return jsonify({"msg": "Document not found"}), 404
    if not doc.requires_return and doc.action_type not in ('sign', 'fill', 'sign_fill'):
        return jsonify({"msg": "This document does not require a return upload"}), 400

    if 'file' not in request.files:
        return jsonify({"msg": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"msg": "No selected file"}), 400
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > 15 * 1024 * 1024:
        return jsonify({"msg": "File too large (max 15MB)"}), 400

    ext = '.' + file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else '.pdf'
    if ext not in ['.pdf', '.doc', '.docx', '.png', '.jpg', '.jpeg']:
        return jsonify({"msg": "PDF, Word, or image files only"}), 400
    filename = safe_filename('return', file.filename)
    file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
    file_url = f"/api/uploads/{filename}"

    history = []
    if doc.returned_files_history:
        try:
            history = json.loads(doc.returned_files_history)
        except:
            history = []
    if doc.returned_file_path:
        history.append({
            'file_path': doc.returned_file_path,
            'uploaded_at': doc.updated_at.isoformat() if doc.updated_at else datetime.now(timezone.utc).isoformat(),
            'status_before': doc.status
        })
    doc.returned_files_history = json.dumps(history, ensure_ascii=False)
    doc.returned_file_path = file_url
    # Admin-initiated: upload at PENDING_REVIEW or AWAITING_RETURN → AWAITING_ADMIN
    # Otherwise: standard return or revision re-upload → RETURNED
    if doc.requested_by == 'ADMIN' and doc.status in ('PENDING_REVIEW', 'AWAITING_RETURN'):
        doc.status = 'AWAITING_ADMIN'
    else:
        doc.status = 'RETURNED'
    doc.updated_at = datetime.now(timezone.utc)

    # Notify Admin and Manager about the return
    doc_title = doc.custom_title or doc.doc_type
    notif_type = 'GENERIC' if doc.status == 'AWAITING_ADMIN' else 'DOCUMENT_RETURNED'
    notif_body = f'تم رفع النسخة المعبأة من "{doc_title}" من قبل المتدرب بانتظار قبولك.' if doc.status == 'AWAITING_ADMIN' else f'تم إرجاع "{doc_title}" من قبل المتدرب وهو قيد المراجعة.'
    notif = Notification(
        intern_id=intern_id, type=notif_type,
        title=f'إعادة مستند من {intern.name if intern else "متدرب"}',
        body=notif_body,
        related_doc_id=doc.id
    )
    db.session.add(notif)
    db.session.commit()

    return jsonify({"msg": "تم استلام النسخة المعبأة", "status": doc.status, "doc_id": doc.id}), 200


@app.route('/api/interns/<int:intern_id>/documents/<int:doc_id>/admin-upload', methods=['POST'])
@jwt_required()
@check_permission('doc_reupload', 'add')
def admin_upload_signed(intern_id, doc_id):
    """Admin uploads a signed/filled version in response to an intern's sign/fill request."""
    current_user = get_jwt()
    doc = DocumentLifecycle.query.filter_by(id=doc_id, intern_id=intern_id).first()
    if not doc:
        return jsonify({"msg": "Document not found"}), 404
    if doc.requested_by == 'ADMIN' and doc.status not in ('REVISION_REQUESTED', 'AWAITING_ADMIN', 'RETURNED'):
        return jsonify({"msg": "Not an intern-initiated request"}), 400
    if doc.action_type not in ('sign', 'fill', 'sign_fill', 'view', 'view_or_return'):
        return jsonify({"msg": "Document action type does not require admin upload"}), 400

    if 'file' not in request.files:
        return jsonify({"msg": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"msg": "No selected file"}), 400

    filename = safe_filename('admin_resp', file.filename)
    file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
    file_url = f"/api/uploads/{filename}"

    history = []
    if doc.returned_files_history:
        try:
            history = json.loads(doc.returned_files_history)
        except:
            history = []
    if doc.returned_file_path:
        history.append({
            'file_path': doc.returned_file_path,
            'uploaded_at': doc.updated_at.isoformat() if doc.updated_at else datetime.now(timezone.utc).isoformat(),
            'uploaded_by': 'ADMIN',
            'status_before': doc.status
        })
    doc.returned_files_history = json.dumps(history, ensure_ascii=False)
    doc.returned_file_path = file_url
    if doc.action_type == 'view':
        doc.status = 'APPROVED_AND_SIGNED'
    elif doc.requested_by == 'ADMIN' and doc.status == 'REVISION_REQUESTED':
        doc.status = 'AWAITING_RETURN'
    else:
        doc.status = 'AWAITING_INTERN'
    doc.updated_at = datetime.now(timezone.utc)

    intern = db.session.get(Intern, intern_id)
    doc_title = doc.custom_title or doc.doc_type
    notif = Notification(
        intern_id=intern_id, type='GENERIC',
        title=f'تم رفع {doc_title} من الإدارة',
        body=f'تم رفع نسخة موقعة/معبأة من "{doc_title}" من قبل الإدارة بانتظار قبولك.',
        related_doc_id=doc.id
    )
    db.session.add(notif)
    db.session.commit()

    log_action(current_user.get('name'), f"رفع نسخة موقعة/معبأة من {doc_title} للمتدرب {intern.name if intern else ''}")
    return jsonify({"msg": "تم رفع النسخة بنجاح", "status": doc.status, "doc_id": doc.id}), 200


@app.route('/api/interns/<int:intern_id>/documents/<int:doc_id>/admin-delete-upload', methods=['POST'])
@jwt_required()
@check_permission('doc_reupload', 'delete')
def admin_delete_upload(intern_id, doc_id):
    """Admin deletes only their uploaded signed/filled version (keeps the request)."""
    current_user = get_jwt()
    doc = DocumentLifecycle.query.filter_by(id=doc_id, intern_id=intern_id).first()
    if not doc:
        return jsonify({"msg": "Document not found"}), 404
    if not doc.returned_file_path:
        return jsonify({"msg": "No admin upload to delete"}), 400

    doc.returned_file_path = None
    doc.returned_files_history = None
    doc.status = 'PENDING_REVIEW'
    doc.updated_at = datetime.now(timezone.utc)
    db.session.commit()

    log_action(current_user.get('name'), f"حذف الرفع الإداري للمستند {doc.custom_title or doc.doc_type}")
    return jsonify({"msg": "تم حذف النسخة المرفوعة"}), 200


@app.route('/api/interns/<int:intern_id>/documents/<int:doc_id>/accept-request', methods=['POST'])
@jwt_required()
@check_permission('doc_reupload', 'approve')
def admin_accept_request(intern_id, doc_id):
    """Admin accepts the intern's sign/fill request, marking it as in progress."""
    current_user = get_jwt()
    doc = DocumentLifecycle.query.filter_by(id=doc_id, intern_id=intern_id).first()
    if not doc:
        return jsonify({"msg": "Document not found"}), 404
    if doc.status != 'AWAITING_ADMIN':
        return jsonify({"msg": "Document is not awaiting admin action"}), 400

    doc.status = 'PENDING_REVIEW'
    doc.updated_at = datetime.now(timezone.utc)
    db.session.commit()

    log_action(current_user.get('name'), f"قبول طلب توقيع/تعبئة للمستند {doc.custom_title or doc.doc_type}")
    return jsonify({"msg": "تم قبول الطلب"}), 200


@app.route('/api/interns/<int:intern_id>/documents/<int:doc_id>/intern-accept-request', methods=['POST'])
@jwt_required()
def intern_accept_request(intern_id, doc_id):
    """Intern accepts the admin's sign/fill request, marking it as in progress."""
    intern, claims, user = _get_doc_type_intern()
    if not intern or intern.id != intern_id:
        return jsonify({"msg": "Unauthorized"}), 403

    doc = DocumentLifecycle.query.filter_by(id=doc_id, intern_id=intern_id).first()
    if not doc:
        return jsonify({"msg": "Document not found"}), 404
    if doc.status != 'AWAITING_RETURN':
        return jsonify({"msg": "Document is not awaiting return"}), 400
    if doc.requested_by != 'ADMIN':
        return jsonify({"msg": "Not an admin-initiated request"}), 400
    if doc.action_type not in ('sign', 'fill', 'sign_fill'):
        return jsonify({"msg": "Not a sign/fill request"}), 400

    doc.status = 'APPROVED_AND_SIGNED'
    doc.updated_at = datetime.now(timezone.utc)

    # Also accept the paired fill doc for sign_fill groups
    if doc.action_type == 'sign_fill':
        fill_doc = DocumentLifecycle.query.filter_by(
            intern_id=intern_id, doc_group=doc.doc_group, action_type='fill'
        ).first()
        if fill_doc:
            fill_doc.status = 'APPROVED_AND_SIGNED'
            fill_doc.updated_at = datetime.now(timezone.utc)

    db.session.commit()

    log_action(intern.name or 'متدرب', f"قبول طلب توقيع/تعبئة من الإدارة للمستند {doc.custom_title or doc.doc_type}")
    return jsonify({"msg": "تم قبول الطلب"}), 200


@app.route('/api/interns/<int:intern_id>/documents/<int:doc_id>/intern-accept', methods=['POST'])
@jwt_required()
def intern_accept_admin_upload(intern_id, doc_id):
    """Intern accepts the admin's uploaded signed/filled version."""
    intern, claims, user = _get_doc_type_intern()
    if not intern or intern.id != intern_id:
        return jsonify({"msg": "Unauthorized"}), 403

    doc = DocumentLifecycle.query.filter_by(id=doc_id, intern_id=intern_id).first()
    if not doc:
        return jsonify({"msg": "Document not found"}), 404
    if doc.status != 'AWAITING_INTERN':
        return jsonify({"msg": "Document not awaiting intern approval"}), 400

    doc.status = 'APPROVED_AND_SIGNED'
    doc.is_visible_to_intern = True
    doc.updated_at = datetime.now(timezone.utc)
    db.session.commit()

    doc_title = doc.custom_title or doc.doc_type
    admin_notif = Notification(intern_id=intern_id, type='GENERIC',
        title=f'قبول مستند من {intern.name or "متدرب"}',
        body=f'قام المتدرب بقبول "{doc_title}" المرفوع من الإدارة.',
        related_doc_id=doc.id)
    db.session.add(admin_notif)
    db.session.commit()

    log_action(intern.name or 'متدرب', f"قبول المستند {doc_title} من الإدارة")
    return jsonify({"msg": "تم قبول المستند"}), 200


@app.route('/api/interns/<int:intern_id>/documents/<int:doc_id>/intern-request-revision', methods=['POST'])
@jwt_required()
def intern_request_revision(intern_id, doc_id):
    """Intern requests admin to re-upload a corrected version."""
    intern, claims, user = _get_doc_type_intern()
    if not intern or intern.id != intern_id:
        return jsonify({"msg": "Unauthorized"}), 403

    doc = DocumentLifecycle.query.filter_by(id=doc_id, intern_id=intern_id).first()
    if not doc:
        return jsonify({"msg": "Document not found"}), 404
    if doc.status not in ('AWAITING_INTERN', 'PENDING_REVIEW', 'AWAITING_ADMIN', 'AWAITING_RETURN'):
        return jsonify({"msg": "Document not in a revisable state"}), 400

    data = request.json or {}
    reason = (data.get('reason') or '').strip()

    doc.status = 'REVISION_REQUESTED'
    doc.rejection_reason = reason
    doc.revision_requested_by = 'INTERN'
    doc.updated_at = datetime.now(timezone.utc)
    db.session.commit()

    doc_title = doc.custom_title or doc.doc_type
    admin_notif = Notification(intern_id=intern_id, type='GENERIC',
        title=f'طلب إعادة رفع من {intern.name or "متدرب"}',
        body=f'طلب المتدرب إعادة رفع "{doc_title}"{f" بسبب: {reason}" if reason else ""}.',
        related_doc_id=doc.id)
    db.session.add(admin_notif)
    db.session.commit()

    log_action(intern.name or 'متدرب', f"طلب إعادة رفع المستند {doc_title}")
    return jsonify({"msg": "تم طلب إعادة الرفع"}), 200


@app.route('/api/intern-documents/<int:doc_id>/download', methods=['GET'])
def download_intern_document(doc_id):
    """Secure download: intern can only download if is_visible_to_intern or they uploaded it."""
    doc = db.session.get(DocumentLifecycle, doc_id)
    if not doc or not doc.file_path:
        return jsonify({"msg": "Document not found"}), 404

    # Determine requester role
    token = request.args.get('token')
    intern_identity = None
    is_admin = False
    if token:
        try:
            from flask_jwt_extended import decode_token
            decoded = decode_token(token)
            role = decoded.get('role')
            is_admin = role in ('Admin', 'Manager')
            if role == 'Intern':
                user = db.session.get(User, int(decoded.get('sub', 0)))
                if user:
                    intern_identity = Intern.query.filter_by(email=user.email).first()
        except Exception:
            import jwt
            try:
                decoded = jwt.decode(token, app.config['JWT_SECRET_KEY'], algorithms=['HS256'], options={"verify_exp": False})
                role = decoded.get('role')
                is_admin = role in ('Admin', 'Manager')
            except Exception as e:
                print(f"Silent Error Caught: {e}")

    if not is_admin:
        if not intern_identity or intern_identity.id != doc.intern_id:
            return jsonify({"msg": "Unauthorized"}), 403
        if intern_identity.status == 'مرفوض':
            return jsonify({"msg": "Account not yet activated"}), 403
        if not doc.is_visible_to_intern and doc.uploaded_by == 'ADMIN':
            return jsonify({"msg": "Unauthorized"}), 403

    is_returned = request.args.get('returned') == '1'
    file_path = doc.returned_file_path if (is_returned and doc.returned_file_path) else doc.file_path
    if not file_path:
        return jsonify({"msg": "File not found"}), 404
    name = file_path.replace('/api/uploads/', '').replace('/', '')
    ext = os.path.splitext(name)[1] or '.pdf'
    download_name = (doc.custom_title or doc.doc_type or 'document') + ext
    return send_file(os.path.join(app.config['UPLOAD_FOLDER'], name), download_name=download_name)



@app.route('/api/intern-documents/<int:doc_id>/open', methods=['GET'])
def open_intern_document(doc_id):
    """Open intern document locally using os.startfile."""
    doc = db.session.get(DocumentLifecycle, doc_id)
    if not doc or not doc.file_path:
        return jsonify({"msg": "Document not found"}), 404

    # Determine requester role
    token = request.args.get('token')
    intern_identity = None
    is_admin = False
    if token:
        try:
            from flask_jwt_extended import decode_token
            decoded = decode_token(token)
            role = decoded.get('role')
            is_admin = role in ('Admin', 'Manager')
            if role == 'Intern':
                user = db.session.get(User, int(decoded.get('sub', 0)))
                if user:
                    intern_identity = Intern.query.filter_by(email=user.email).first()
        except Exception as e:
                print(f"Silent Error Caught: {e}")

    if not is_admin:
        if not intern_identity or intern_identity.id != doc.intern_id:
            return jsonify({"msg": "Unauthorized"}), 403
        if not doc.is_visible_to_intern and doc.uploaded_by == 'ADMIN':
            return jsonify({"msg": "Unauthorized"}), 403

    is_returned = request.args.get('returned') == '1'
    file_path = doc.returned_file_path if (is_returned and doc.returned_file_path) else doc.file_path
    if not file_path:
        return jsonify({"msg": "File not found"}), 404
        
    name = file_path.replace('/api/uploads/', '').replace('/', '')
    full_path = os.path.join(app.config['UPLOAD_FOLDER'], name)
    if os.path.exists(full_path):
        import subprocess
        if os.name == 'nt':
            os.startfile(full_path)
        else:
            subprocess.call(['open', full_path])
        return jsonify({"msg": "Opened successfully"}), 200
    return jsonify({"msg": "File not found locally"}), 404


@app.route('/api/intern/documents', methods=['GET'])
@jwt_required()
def list_my_documents():
    """Intern portal: list all documents visible to the intern."""
    intern, claims, user = _get_doc_type_intern()
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404
    if intern.status == 'مرفوض':
        return jsonify({"msg": "Account not yet activated"}), 403
    docs = DocumentLifecycle.query.filter_by(intern_id=intern.id).order_by(DocumentLifecycle.created_at.asc()).all()
    result = []
    for d in docs:
        # The intern only sees rows shared with them or that they uploaded.
        if not (d.is_visible_to_intern or (d.uploaded_by or '') == 'INTERN'):
            continue
        entry = {
            "id": d.id,
            "doc_type": d.doc_type,
            "label": d.custom_title or d.doc_type,
            "file_path": d.file_path,
            "uploaded_by": d.uploaded_by,
            "status": d.status,
            "rejection_reason": d.rejection_reason,
            "is_visible_to_intern": d.is_visible_to_intern,
            "custom_title": d.custom_title,
            "requires_return": d.requires_return,
            "returned_file_path": d.returned_file_path,
            "returned_files_history": d.returned_files_history,
            "created_at": d.created_at,
            "updated_at": d.updated_at,
            "file_type": d.file_type or 'pdf',
            "action_type": d.action_type or 'view',
            "requested_by": d.requested_by or ('INTERN' if (d.uploaded_by or '') == 'INTERN' else 'ADMIN'),
            "revision_requested_by": d.revision_requested_by,
            "source": d.source or 'ADMIN',
        }
        result.append(entry)
    return jsonify(result), 200


# --- ZIP ARCHIVE EXPORT ---


@app.route('/api/interns/<int:intern_id>/export-zip', methods=['GET'])
@jwt_required()
@check_permission('interns', 'approve')
def export_intern_zip(intern_id):
    intern = db.session.get(Intern, intern_id)
    if not intern:
        return jsonify({"msg": "Intern not found"}), 404
    import tempfile, zipfile, json, re
    
    docs = DocumentLifecycle.query.filter_by(intern_id=intern_id).all()
    
    old_docs = {}
    if intern.documents:
        try:
            old_docs = json.loads(intern.documents)
        except Exception as e: print(f"Silent Error Caught: {e}")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.zip')
    with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zf:
        added_files = set()
        
        # Helper to get the correct suffix based on status and action_type
        def get_suffix(d):
            if d.status == 'APPROVED_AND_SIGNED':
                return "_Signed"
            if d.status == 'RETURNED':
                return "_Filled_And_Returned"
            if d.status == 'AWAITING_INTERN':
                act = str(d.action_type or '').lower()
                lc = str(d.lifecycle_type or '').lower()
                req = act if act and act != 'view' else lc
                if 'sign_fill' in req or ('sign' in req and 'fill' in req):
                    return "_Requested_To_Fill_And_Sign"
                elif 'sign' in req:
                    return "_Requested_To_Sign"
                elif 'fill' in req:
                    return "_Requested_To_Fill"
                return "_Requested_To_Review"
            return ""

        # Folders
        dir_admin = "01_Administration_To_Intern"
        dir_intern = "02_Intern_To_Administration"
        dir_history = "03_History_And_Returns"
        dir_all = "04_All_Files"

        # Add old documents (mostly from intern)
        for k, v in old_docs.items():
            if v and isinstance(v, str):
                fname = v.replace('/api/uploads/', '').replace('/', '')
                fpath = os.path.join(app.config['UPLOAD_FOLDER'], fname)
                if os.path.exists(fpath):
                    name_base, ext = os.path.splitext(fname)
                    arc_name = f"{k}{ext}"
                    # Add to Intern to Admin and All
                    path1 = f"{dir_intern}/{arc_name}"
                    path2 = f"{dir_all}/old_{arc_name}"
                    if path1 not in added_files:
                        zf.write(fpath, path1)
                        added_files.add(path1)
                    if path2 not in added_files:
                        zf.write(fpath, path2)
                        added_files.add(path2)
        
        # Add lifecycle documents
        for d in docs:
            suffix = get_suffix(d)
            
            # Original file (uploaded by admin or intern)
            if d.file_path:
                fname = d.file_path.replace('/api/uploads/', '').replace('/', '')
                fpath = os.path.join(app.config['UPLOAD_FOLDER'], fname)
                if os.path.exists(fpath):
                    title = re.sub(r'[^\w\-\.]', '_', d.custom_title or d.doc_type or f"doc_{d.id}")
                    _, ext = os.path.splitext(fname)
                    arc_name = f"{title}{suffix}{ext}"
                    
                    folder = dir_admin if d.uploaded_by == 'ADMIN' else dir_intern
                    path1 = f"{folder}/{arc_name}"
                    path2 = f"{dir_all}/{arc_name}"
                    
                    if path1 not in added_files:
                        zf.write(fpath, path1)
                        added_files.add(path1)
                    if path2 not in added_files:
                        zf.write(fpath, path2)
                        added_files.add(path2)
            
            # Latest Return (from intern)
            if d.returned_file_path:
                fname = d.returned_file_path.replace('/api/uploads/', '').replace('/', '')
                fpath = os.path.join(app.config['UPLOAD_FOLDER'], fname)
                if os.path.exists(fpath):
                    title = re.sub(r'[^\w\-\.]', '_', d.custom_title or d.doc_type or f"doc_{d.id}")
                    _, ext = os.path.splitext(fname)
                    arc_name = f"{title}_Latest_Return{ext}"
                    
                    path1 = f"{dir_intern}/{arc_name}"
                    path2 = f"{dir_all}/{arc_name}"
                    
                    if path1 not in added_files:
                        zf.write(fpath, path1)
                        added_files.add(path1)
                    if path2 not in added_files:
                        zf.write(fpath, path2)
                        added_files.add(path2)

            # History Returns
            if d.returned_files_history:
                try:
                    history = json.loads(d.returned_files_history)
                    for idx, entry in enumerate(history):
                        fname = entry.get('file_path', '').replace('/api/uploads/', '').replace('/', '')
                        if not fname:
                            continue
                        fpath = os.path.join(app.config['UPLOAD_FOLDER'], fname)
                        if os.path.exists(fpath):
                            title = re.sub(r'[^\w\-\.]', '_', d.custom_title or d.doc_type or f"doc_{d.id}")
                            _, ext = os.path.splitext(fname)
                            arc_name = f"{title}_v{idx+1}{ext}"
                            
                            path1 = f"{dir_history}/{arc_name}"
                            path2 = f"{dir_all}/{arc_name}"
                            
                            if path1 not in added_files:
                                zf.write(fpath, path1)
                                added_files.add(path1)
                            if path2 not in added_files:
                                zf.write(fpath, path2)
                                added_files.add(path2)
                except Exception as e: print(f"Silent Error Caught: {e}")

        # Write README
        readme_content = """دليل ملفات المتدرب

يهدف هذا الأرشيف إلى تنظيم ملفات المتدرب بناءً على مصدرها وحالتها:

1. مجلد "01_Administration_To_Intern": 
يحتوي على النماذج الفارغة التي يطلب من المتدرب تحميلها وتعبئتها، بالإضافة إلى الشهادات الموقعة النهائية التي تمنحها الإدارة.

2. مجلد "02_Intern_To_Administration": 
يحتوي على الوثائق الشخصية للمتدرب (مثل السيرة الذاتية والبطاقة الوطنية) والوثائق التي قام المتدرب بتعبئتها وإعادتها.

3. مجلد "03_History_And_Returns": 
يحتوي على النسخ القديمة من الوثائق في حال تم إرجاعها للمتدرب لتعديلها عدة مرات.

4. مجلد "04_All_Files": 
يحتوي على جميع الوثائق والملفات بدون تقسيم للمجلدات.

---

شرح الملحقات المضافة لأسماء الملفات:
- Requested_To_Fill_And_Sign : نموذج طلب من المتدرب تعبئته وتوقيعه.
- Requested_To_Sign : نموذج يطلب من المتدرب توقيعه فقط.
- Requested_To_Fill : نموذج يطلب تعبئته فقط.
- Filled_And_Returned : وثيقة قام المتدرب بتعبئتها وإرسالها للإدارة وهي بانتظار المراجعة والاعتماد.
- Signed : وثيقة نهائية تمت الموافقة عليها وتم توقيعها بنجاح.
"""
        zf.writestr('README.txt', readme_content.encode('utf-8'))

    tmp.close()
    
    # Safe filename
    safe_name = re.sub(r'[^\w\-\s]', '', intern.name or f"Intern_{intern_id}").strip().replace(' ', '_')
    if not safe_name:
        safe_name = f"Intern_{intern_id}"
        
    return send_file(tmp.name, as_attachment=True, download_name=f"{safe_name}_Files.zip", mimetype='application/zip')



# --- FORM BUILDER ROUTES ---

@app.route('/api/forms', methods=['GET'])
@jwt_required()
@check_permission('forms', 'view')
def get_forms():
    forms = Form.query.order_by(Form.id.desc()).all()
    return jsonify([{
        "id": f.id, "title": f.title or "نموذج بدون عنوان",
        "slug": f.slug, "is_active": f.is_active,
        "created_at": f.created_at,
        "form_data": f.form_data,
        "pending_count": FormSubmission.query.filter_by(form_id=f.id, status='pending').count()
    } for f in forms])

@app.route('/api/forms', methods=['POST'])
@jwt_required()
@check_permission('forms', 'add')
def save_form():
    data = request.json
    fields = data.get('form_data', [])
    title = data.get('title', 'نموذج تسجيل')
    form_id = data.get('id')  # if editing existing form
    now = datetime.now(timezone.utc).isoformat()

    if form_id:
        form = db.session.get(Form, form_id)
        if not form:
            return jsonify({"msg": "النموذج غير موجود"}), 404
        form.form_data = json.dumps(fields, ensure_ascii=False)
        form.title = title
    else:
        slug = uuid.uuid4().hex[:8]
        form = Form(form_data=json.dumps(fields, ensure_ascii=False), title=title, slug=slug, is_active=True, created_at=now)
        db.session.add(form)

    db.session.commit()
    return jsonify({"success": True, "id": form.id, "slug": form.slug, "title": form.title})

@app.route('/api/forms/<int:form_id>', methods=['DELETE'])
@jwt_required()
@check_permission('forms', 'delete')
def delete_form(form_id):
    form = db.session.get(Form, form_id)
    if not form:
        return jsonify({"msg": "النموذج غير موجود"}), 404
    db.session.delete(form)
    db.session.commit()
    return jsonify({"success": True})

@app.route('/api/forms/<int:form_id>/toggle', methods=['POST'])
@jwt_required()
@check_permission('forms', 'edit')
def toggle_form(form_id):
    form = db.session.get(Form, form_id)
    if not form:
        return jsonify({"msg": "النموذج غير موجود"}), 404
    form.is_active = not form.is_active
    db.session.commit()
    return jsonify({"success": True, "is_active": form.is_active})

# --- PUBLIC FORM ROUTES (no JWT) ---

@app.route('/api/public-form/<slug>', methods=['GET'])
def get_public_form(slug):
    form = Form.query.filter_by(slug=slug).first()
    if not form:
        return jsonify({"success": False, "msg": "النموذج غير موجود"}), 404
    if not form.is_active:
        return jsonify({"success": False, "msg": "النموذج غير متاح حالياً"}), 403
    try:
        fields = json.loads(form.form_data)
    except Exception:
        fields = []
    return jsonify({"success": True, "form_data": fields, "title": form.title or "استمارة التقديم للتدريب"})

@app.route('/api/public-form/<slug>/submit', methods=['POST'])
def submit_public_form(slug):
    form = Form.query.filter_by(slug=slug).first()
    if not form or not form.is_active:
        return jsonify({"success": False, "msg": "النموذج غير متاح"}), 403

    data = request.json or {}
    now = datetime.now(timezone.utc).isoformat()

    submission = FormSubmission(
        form_id=form.id,
        submitted_data=json.dumps(data, ensure_ascii=False),
        status='pending',
        submitted_at=now
    )
    db.session.add(submission)
    db.session.flush()
    
    # Auto convert to Intern
    new_intern = _process_submission_to_intern(submission)
    new_intern.source = form.title or 'نموذج محلي'
    db.session.commit()

    # Send confirmation email to intern if email field exists
    intern_email = None
    try:
        fields = json.loads(form.form_data)
        for field in fields:
            if field.get('maps_to') == 'email' and data.get(field['label']):
                intern_email = data[field['label']]
                break
        # fallback: look for any email-type field
        if not intern_email:
            for field in fields:
                if field.get('type') == 'email' and data.get(field['label']):
                    intern_email = data[field['label']]
                    break
    except Exception as e:
                print(f"Silent Error Caught: {e}")

    if intern_email:
        email_service.send_received_email(intern_email, data.get('الاسم', 'متدرب'))

    return jsonify({"success": True, "msg": "تم إرسال طلبك بنجاح"})

@app.route('/api/public-upload', methods=['POST'])
def public_upload():
    if 'file' not in request.files:
        return jsonify({"msg": "لا يوجد ملف"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"msg": "لا يوجد ملف محدد"}), 400
    filename = safe_filename('pub', file.filename)
    file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
    return jsonify({"success": True, "filename": filename, "path": f"/api/uploads/{filename}"})

# --- ADMIN: SUBMISSIONS ROUTES ---

@app.route('/api/submissions', methods=['GET'])
@jwt_required()
@check_permission('approve_interns', 'view')
def get_submissions():
    status_filter = request.args.get('status', 'pending')
    submissions = FormSubmission.query.filter_by(status=status_filter).order_by(FormSubmission.id.desc()).all()
    result = []
    for s in submissions:
        try:
            data = json.loads(s.submitted_data)
        except Exception:
            data = {}
        form = db.session.get(Form, s.form_id)
        result.append({
            "id": s.id,
            "form_id": s.form_id,
            "form_title": form.title if form else "—",
            "submitted_data": data,
            "source": data.get('_source', 'إضافة يدوية'),
            "status": s.status,
            "submitted_at": s.submitted_at,
            "rejection_reason": s.rejection_reason,
            "intern_id": s.intern_id
        })
    return jsonify(result)

def _map_submission_fields(data, mapping=None):
    """Turn a raw submission row (Arabic titles or canonical keys) into an intern field dict.

    Uses the same exact-mapping + generic fallback logic for both Google Form rows
    and form-builder submissions.
    """
    mapping = mapping or {}

    # Add exact mapping for the manual Google Forms guide
    exact_mapping = {
        "الإسم الكامل (بالعربية)": "name",
        "الإسم الكامل (بالفرنسة)": "name_fr",
        "البريد الإلكتروني": "email",
        "رقم البطاقة الوطنية": "national_id",
        "رقم الهاتف": "phone",
        "الجامعة / المؤسسة": "university",
        "التخصص": "specialty",
        "تاريخ بدء التدريب": "start_date",
        "تاريخ إنتهاء التدريب": "end_date",
        "تاريخ الميلاد": "date_of_birth",
        "العنوان": "address",
        "القسم / الدائرة": "department",
        "الصورة الشخصية": "photo_path"
    }

    intern_data = {'name': None, 'name_fr': None, 'email': None, 'national_id': None,
                   'phone': None, 'university': None, 'specialty': None, 'start_date': None, 'end_date': None,
                   'date_of_birth': None, 'address': None, 'department': None, 'photo_path': None}

    for label, value in data.items():
        mapped_field = mapping.get(label) or exact_mapping.get(label.strip())
        if mapped_field and mapped_field in intern_data:
            intern_data[mapped_field] = value

    # Use generic fallbacks if mapping missed something
    for k, v in data.items():
        if k == '_source': continue
        kl = str(k).lower().replace('إ', 'ا').replace('أ', 'ا').replace('آ', 'ا').replace('ة', 'ه')
        if not intern_data['name'] and ('name' in kl or ('اسم' in kl and 'عرب' in kl)):
            intern_data['name'] = v
        elif not intern_data['name'] and ('اسم' in kl or 'name' in kl):
            intern_data['name'] = v

        if not intern_data['name_fr'] and ('name_fr' in kl or 'فرنس' in kl or 'nom' in kl):
            intern_data['name_fr'] = v
        if not intern_data['email'] and ('email' in kl or 'بريد' in kl):
            intern_data['email'] = v
        if not intern_data['national_id'] and ('national' in kl or 'بطاق' in kl or 'cin' in kl):
            intern_data['national_id'] = v
        if not intern_data['phone'] and ('phone' in kl or 'هاتف' in kl or 'tel' in kl):
            intern_data['phone'] = v
        if not intern_data['university'] and ('university' in kl or 'جامع' in kl or 'مؤسس' in kl or 'معهد' in kl):
            intern_data['university'] = v
        if not intern_data['specialty'] and ('specialty' in kl or 'تخصص' in kl or 'شعب' in kl):
            intern_data['specialty'] = v
        if not intern_data['start_date'] and ('start' in kl or 'بدء' in kl or 'بداي' in kl):
            intern_data['start_date'] = v
        if not intern_data['end_date'] and ('end' in kl or 'نهاي' in kl or 'إنتهاء' in kl or 'انتهاء' in kl):
            intern_data['end_date'] = v
        if not intern_data['date_of_birth'] and ('birth' in kl or 'ميلاد' in kl or 'ازدياد' in kl):
            intern_data['date_of_birth'] = v

    if not intern_data['name']:
        intern_data['name'] = 'متدرب جديد'

    # Fallback for photo_path
    if not intern_data['photo_path']:
        for k, v in data.items():
            if 'photo' in k.lower() or 'صورة' in k:
                intern_data['photo_path'] = v
                break

    # Convert Google Drive URL to direct image link
    photo_path = intern_data.get('photo_path')
    if photo_path and isinstance(photo_path, str) and 'drive.google.com/open?id=' in photo_path:
        intern_data['photo_path'] = photo_path.replace('open?id=', 'uc?export=view&id=')

    return intern_data


def create_intern_record(data, source):
    """Create an Intern using the SAME code as the manual add (add_intern).

    Shared by the manual 'إضافة يدوية' route and the Google Form sync so both
    write to the same table/fields, differing only in `source`.
    """
    us_format = source == 'نمازج جوجل'
    new_intern = Intern(
        name=data.get('name') or 'متدرب جديد',
        name_fr=data.get('name_fr'),
        email=data.get('email'),
        national_id=data.get('national_id'),
        department=data.get('department'),
        encadrant=data.get('encadrant'),
        phone=data.get('phone'),
        start_date=_parse_date(data.get('start_date'), us_format=us_format),
        end_date=_parse_date(data.get('end_date'), us_format=us_format),
        date_of_birth=_parse_date(data.get('date_of_birth'), us_format=us_format),
        university=data.get('university'),
        specialty=data.get('specialty'),
        address=data.get('address'),
        photo_path=data.get('photo_path'),
        status=data.get('status', 'قيد المراجعة'),
        source=source,
        documents=json.dumps(data.get('documents', {}))
    )
    db.session.add(new_intern)
    db.session.flush()

    # Create user account automatically
    create_user_for_intern(new_intern)

    # Auto-create DocumentLifecycle entries for active templates
    templates = DocumentTemplate.query.filter_by(is_active=True).all()
    now = datetime.now(timezone.utc)
    for t in templates:
        record = DocumentLifecycle(
            intern_id=new_intern.id, doc_type='OTHER', status='MISSING',
            uploaded_by='ADMIN', is_visible_to_intern=True,
            custom_title=t.label, action_type='view',
            created_at=now, updated_at=now,
            source='TEMPLATE_VIEW'
        )
        db.session.add(record)

    if new_intern.status == 'قيد المراجعة' and new_intern.email:
        try:
            email_service.send_received_email(new_intern.email, new_intern.name)
        except Exception as e:
            print("Error sending received email:", e)

    return new_intern


def _process_submission_to_intern(submission):
    try:
        data = json.loads(submission.submitted_data)
    except Exception:
        data = {}

    form = db.session.get(Form, submission.form_id)
    mapping = {}
    if form:
        try:
            fields = json.loads(form.form_data)
            mapping = {f['label']: f.get('maps_to') for f in fields if f.get('maps_to')}
        except Exception as e:
                print(f"Silent Error Caught: {e}")

    intern_data = _map_submission_fields(data, mapping)
    return create_intern_record(intern_data, 'نمازج جوجل')

@app.route('/api/submissions/<int:sub_id>/approve', methods=['POST'])
@jwt_required()
@check_permission('approve_interns', 'approve')
def approve_submission(sub_id):
    current_user = get_jwt()
    submission = db.session.get(FormSubmission, sub_id)
    if not submission:
        return jsonify({"msg": "الطلب غير موجود"}), 404
    if submission.status != 'pending':
        return jsonify({"msg": "الطلب ليس في حالة معلقة"}), 400

    new_intern = _process_submission_to_intern(submission)
    
    submission.status = 'approved'
    submission.intern_id = new_intern.id
    db.session.commit()

    log_action(get_jwt_identity(), f"قبول طلب تسجيل #{sub_id} وإنشاء ملف متدرب #{new_intern.id}")

    # Send approval email
    try:
        data = json.loads(submission.submitted_data)
        intern_email = None
        for key, val in data.items():
            if 'email' in key.lower() or 'بريد' in key:
                intern_email = val
                break
        if intern_email:
            email_service.send_accepted_email(intern_email, new_intern.name or 'متدرب')
    except Exception as e:
                print(f"Silent Error Caught: {e}")

    return jsonify({"success": True, "intern_id": new_intern.id})

@app.route('/api/submissions/<int:sub_id>/reject', methods=['POST'])
@jwt_required()
@check_permission('approve_interns', 'approve')
def reject_submission(sub_id):
    current_user = get_jwt()
    submission = db.session.get(FormSubmission, sub_id)
    if not submission:
        return jsonify({"msg": "الطلب غير موجود"}), 404
    if submission.status != 'pending':
        return jsonify({"msg": "الطلب تمت معالجته بالفعل"}), 400

    reason = request.json.get('reason', '') if request.json else ''
    submission.status = 'rejected'
    submission.rejection_reason = reason
    db.session.commit()

    log_action(get_jwt_identity(), f"رفض طلب تسجيل #{sub_id}")

    # Send rejection email
    try:
        data = json.loads(submission.submitted_data)
        form = db.session.get(Form, submission.form_id)
        fields = json.loads(form.form_data) if form else []
        intern_email = None
        for field in fields:
            if field.get('maps_to') == 'email' and data.get(field['label']):
                intern_email = data[field['label']]
                break
            if field.get('type') == 'email' and data.get(field['label']) and not intern_email:
                intern_email = data[field['label']]
        if intern_email:
            email_service.send_rejected_email(intern_email, data.get('name', 'متدرب'), reason)
    except Exception as e:
        print(f"[REJECT EMAIL ERROR] {e}")

    return jsonify({"success": True})


# --- INTEGRATION API ---
@app.route('/api/integration/settings', methods=['GET'])
@jwt_required()
@check_permission('system_settings', 'view')
def get_integration_settings():
    return jsonify(email_service.get_settings())

@app.route('/api/integration/settings', methods=['POST'])
@jwt_required()
@check_permission('system_settings', 'edit')
def save_integration_settings():
    current_user = get_jwt()
    data = request.json
    email_service.save_settings(data)
    log_action(current_user.get('name', 'Admin'), "قام بتحديث إعدادات الربط مع Google Forms & Gmail")
    return jsonify({"success": True})

# --- GOOGLE SHEET ROW NORMALIZATION ---
# Used when a manually-created sheet has degenerate headers (e.g. all columns
# named "Timestamp"): we fall back to mapping columns by position.
# Column order matches the manual Google Sheets guide:
#   col0 Timestamp (skipped), col1 name, col2 name_fr, col3 email,
#   col4 national_id, col5 date_of_birth, col6 phone, col7 university,
#   col8 specialty, col9 start_date, col10 end_date
GOOGLE_COLUMN_FALLBACK = {
    1: 'name', 2: 'name_fr', 3: 'email', 4: 'national_id', 5: 'date_of_birth',
    6: 'phone', 7: 'university', 8: 'specialty', 9: 'start_date', 10: 'end_date',
}

def _norm_key(k):
    return str(k).lower().replace('إ', 'ا').replace('أ', 'ا').replace('آ', 'ا').replace('ة', 'ه')

def extract_email_from_row(row):
    email = ''
    for k, v in row.items():
        if 'email' in k.lower() or 'بريد' in k:
            email = v or ''
            break
    return str(email).strip()

def has_meaningful_data(row):
    """Skip rows with no name and no email (e.g. empty/timestamp-only rows)."""
    values = [str(v or '').strip() for k, v in row.items() if k != '_source']
    if not any(values):
        return False
    email = extract_email_from_row(row)
    if email:
        return True
    for k, v in row.items():
        kl = _norm_key(k)
        if ('اسم' in kl or 'name' in kl) and str(v or '').strip():
            return True
    return False

def normalize_google_rows(res):
    """Convert fetched Google Sheet rows into usable dicts.

    If headers are degenerate (all identical or all empty), rebuild the rows
    positionally using GOOGLE_COLUMN_FALLBACK.
    """
    rows = res.get('data') or []
    headers = [str(h or '').strip() for h in (res.get('headers') or [])]
    raw = res.get('raw') or []

    # Detect degenerate headers: all cells equal (e.g. "Timestamp" repeated) or all empty.
    non_empty_headers = [h for h in headers if h]
    degenerate = (not non_empty_headers) or len(set(non_empty_headers)) == 1

    if not degenerate:
        return rows

    normalized = []
    for raw_row in raw:
        row = {}
        for idx, value in enumerate(raw_row):
            field = GOOGLE_COLUMN_FALLBACK.get(idx)
            if field:
                row[field] = value
        if row:
            normalized.append(row)
    return normalized

@app.route('/api/forms/sync-google', methods=['POST'])
@jwt_required()
@check_permission('forms', 'add')
def sync_google_forms():
    settings = email_service.get_settings()
    sheet_link = settings.get('google_sheet_link')
    if not sheet_link:
        return jsonify({"success": False, "msg": "يرجى إضافة رابط Google Sheet في الإعدادات أولاً"})
        
    res = google_sheets_service.fetch_google_form_responses(sheet_link)
    if not res['success']:
        return jsonify(res), 400
        
    rows = normalize_google_rows(res)
    added_count = 0
    # To avoid duplicates, skip if an intern with this email already exists.
    for row in rows:
        if not has_meaningful_data(row):
            continue
        email = extract_email_from_row(row)
        if email:
            if Intern.query.filter_by(email=email).first() or BlacklistedIntern.query.filter_by(identifier=email).first():
                continue
        intern_data = _map_submission_fields(row)
        phone = intern_data.get('phone')
        if phone and BlacklistedIntern.query.filter_by(identifier=phone).first():
            continue
        new_intern = create_intern_record(intern_data, 'نمازج جوجل')
        db.session.commit()

        log_action(get_jwt_identity(), f"تم استيراد متدرب جديد من نماذج جوجل: {new_intern.name}")

        added_count += 1
    if added_count > 0:
        log_action(get_jwt_identity(), f"تم مزامنة {added_count} متدرب جديد من Google Forms")

    return jsonify({"success": True, "added": added_count})

# --- BACKGROUND AUTO-SYNC LOOP ---
import threading
import time

def auto_sync_loop():
    while True:
        time.sleep(60)
        try:
            with app.app_context():
                # Google Sync
                try:
                    settings = email_service.get_settings()
                    sheet_link = settings.get('google_sheet_link')
                    if sheet_link:
                        res = google_sheets_service.fetch_google_form_responses(sheet_link)
                        if res['success']:
                            rows = normalize_google_rows(res)
                            added_count = 0
                            for row in rows:
                                if not has_meaningful_data(row):
                                    continue
                                email = extract_email_from_row(row)
                                if email:
                                    if Intern.query.filter_by(email=email).first() or BlacklistedIntern.query.filter_by(identifier=email).first():
                                        continue
                                intern_data = _map_submission_fields(row)
                                phone = intern_data.get('phone')
                                if phone and BlacklistedIntern.query.filter_by(identifier=phone).first():
                                    continue
                                new_intern = create_intern_record(intern_data, 'نمازج جوجل')
                                db.session.commit()

                                added_count += 1
                            if added_count > 0:
                                db.session.commit()
                                log_action("النظام", f"تم مزامنة {added_count} متدرب جديد من Google Forms تلقائيا")
                except Exception as e:
                    print(f"Auto-Sync Google Error: {e}")

                # Microsoft Sync
                try:
                    import microsoft_excel_service
                    settings = email_service.get_settings()
                    excel_link = settings.get('microsoft_excel_link')
                    if excel_link:
                        res = microsoft_excel_service.fetch_microsoft_excel_responses(excel_link)
                        if res['success']:
                            rows = res['data']
                            added_count = 0
                            for row in rows:
                                row['_source'] = 'نماذج مايكروسوفت'
                                row_json = json.dumps(row, ensure_ascii=False)
                                
                                # Duplicate check
                                existing = FormSubmission.query.all()
                                is_duplicate = False
                                for ex in existing:
                                    try:
                                        if json.loads(ex.submitted_data) == row:
                                            is_duplicate = True
                                            break
                                    except Exception as e: print(f"Silent Error Caught: {e}")
                                
                                if not is_duplicate:
                                    default_form = Form.query.first()
                                    new_sub = FormSubmission(
                                        form_id=default_form.id if default_form else 1,
                                        submitted_data=row_json,
                                        status='pending'
                                    )
                                    db.session.add(new_sub)
                                    db.session.flush()

                                    added_count += 1
                            
                            if added_count > 0:
                                db.session.commit()
                                log_action("النظام", f"تم مزامنة {added_count} طلب جديد من Microsoft Forms تلقائيا")
                except Exception as e:
                    print(f"Auto-Sync Microsoft Error: {e}")

        except Exception as e:
            print(f"Fatal Auto-Sync Error: {e}")

# Start background sync thread
threading.Thread(target=auto_sync_loop, daemon=True).start()


@app.route('/api/interns/<int:intern_id>/documents/<int:doc_id>/cancel-request', methods=['POST'])
@jwt_required()
def cancel_intern_request(intern_id, doc_id):
    """Intern cancels their own sign/fill request before admin accepts."""
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"msg": "Unauthorized"}), 401

    intern = Intern.query.filter_by(email=user.email).first()
    if not intern:
        return jsonify({"msg": "Unauthorized"}), 401

    doc = db.session.get(DocumentLifecycle, doc_id)
    if not doc or doc.intern_id != intern.id:
        return jsonify({"msg": "Document not found"}), 400
    if doc.requested_by != 'INTERN':
        return jsonify({"msg": "Not an intern-initiated request"}), 400
    if doc.status != 'AWAITING_ADMIN':
        return jsonify({"msg": "Can only cancel while awaiting admin"}), 400

    Notification.query.filter_by(related_doc_id=doc.id).delete()
    AuditEvent.query.filter_by(doc_id=doc.id).delete()
    DocumentRequest.query.filter_by(lifecycle_id=doc.id).delete()
    DocumentLifecycle.query.filter_by(parent_id=doc.id).delete()
    db.session.delete(doc)
    db.session.commit()
    log_action(user.name, f"ألغى المتدرب طلب {(doc.custom_title or doc.doc_type)}")
    return jsonify({"success": True}), 200



import threading
import time
from datetime import datetime
import traceback

def auto_complete_job():
    while True:
        try:
            with app.app_context():
                today = datetime.now().date()
                interns = Intern.query.all()
                for intern in interns:
                    # 'U.UOU.U,' and 'U.OU?U^O ' (completed and disabled)
                    if intern.end_date and intern.status != 'U.UOU.U,' and intern.status != 'U.O1U?U^O':
                        try:
                            # end_date may be a date object or a dd/mm/yyyy string
                            end_val = intern.end_date
                            if isinstance(end_val, str):
                                end_date_obj = datetime.strptime(end_val, '%d/%m/%Y').date()
                            else:
                                end_date_obj = end_val
                            if today > end_date_obj:
                                # Check if documents are ready
                                required_types = ['CIN', 'CV', 'DEMANDE', 'CONVENTION_SIGNED', 'FINAL_REPORT']
                                docs = DocumentLifecycle.query.filter_by(intern_id=intern.id).all()
                                ready = True
                                for req in required_types:
                                    d = next((x for x in docs if x.doc_type == req), None)
                                    if not d or d.status != 'APPROVED_AND_SIGNED':
                                        ready = False
                                        break
                                
                                if ready:
                                    print(f"Auto-completing internship for intern {intern.id}")
                                    
                                    # Copy of complete_stage logic
                                    notif = Notification(intern_id=intern.id, type='STAGE_COMPLETE', title='تهانينا!',
                                                         body='تم إنهاء فترة تدريبك بنجاح.')
                                    db.session.add(notif)
                                    
                                    intern.status = 'U.UOU.U,' # Maktamal
                                    db.session.commit()
                                    
                                    try:
                                        from reportlab.pdfgen import canvas
                                        from reportlab.lib.pagesizes import A4
                                        import io
                                        buffer = io.BytesIO()
                                        c = canvas.Canvas(buffer, pagesize=A4)
                                        width, height = A4
                                        c.setFont("Helvetica-Bold", 24)
                                        c.drawCentredString(width/2.0, height - 100, "ATTESTATION DE STAGE")
                                        c.setFont("Helvetica", 14)
                                        c.drawString(50, height - 200, f"Nous soussignes, certifions que Monsieur/Madame:")
                                        c.setFont("Helvetica-Bold", 16)
                                        c.drawString(50, height - 230, f"{intern.name_fr or '__________________'}")
                                        c.setFont("Helvetica", 14)
                                        c.drawString(50, height - 280, f"A effectue un stage au sein de notre departement:")
                                        c.setFont("Helvetica-Bold", 14)
                                        c.drawString(50, height - 310, f"{intern.department or '__________________'}")
                                        c.setFont("Helvetica", 14)
                                        c.drawString(50, height - 360, f"Encadre(e) par:")
                                        c.setFont("Helvetica-Bold", 14)
                                        c.drawString(50, height - 390, f"{intern.encadrant or '__________________'}")
                                        c.setFont("Helvetica", 14)
                                        c.drawString(50, height - 440, f"PAcriode: du {intern.start_date or '___'} au {intern.end_date or '___'}")
                                        c.drawString(50, height - 500, "Cette attestation est delivree pour servir et valoir ce que de droit.")
                                        c.drawString(width - 200, 150, "Signature et Cachet:")
                                        c.showPage()
                                        c.save()
                                        
                                        buffer.seek(0)
                                        filename = f"Attestation_Signed_{intern.id}.pdf"
                                        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                                        with open(filepath, 'wb') as f:
                                            f.write(buffer.read())
                                            
                                        file_url = f"/api/uploads/{filename}"
                                        dl = DocumentLifecycle(
                                            intern_id=intern.id, doc_type='ATTESTATION_SIGNED',
                                            file_path=file_url, uploaded_by='ADMIN',
                                            status='APPROVED_AND_SIGNED', is_visible_to_intern=True,
                                            custom_title='شهادة التدريب'
                                        )
                                        db.session.add(dl)
                                        db.session.commit()
                                    except Exception as pdf_err:
                                        print(f"Failed to generate certificate for {intern.id}: {pdf_err}")
                        except Exception as e:
                            print(f"Error checking intern {intern.id}: {e}")
        except Exception as e:
            print(f"Auto complete job error: {traceback.format_exc()}")
        time.sleep(3600 * 24) # Run once a day

# Start thread
auto_complete_thread = threading.Thread(target=auto_complete_job, daemon=True)
auto_complete_thread.start()

# --- Serve React App ---
@app.route('/', defaults={'path': ''})
@app.route('/<path:path>')
def serve_react_app(path):
    if path != "" and os.path.exists(os.path.join(app.static_folder, path)):
        return send_from_directory(app.static_folder, path)
    else:
        return send_from_directory(app.static_folder, 'index.html')

def setup_app():
    init_db()
    # Fix: records uploaded before the status-update bug was fixed
    with app.app_context():
        fixed = DocumentLifecycle.query.filter(
            DocumentLifecycle.file_path.isnot(None),
            DocumentLifecycle.file_path != '',
            DocumentLifecycle.status == 'MISSING'
        ).all()
        for rec in fixed:
            rec.status = 'PENDING_REVIEW'
            rec.updated_at = datetime.now(timezone.utc)
        if fixed:
            db.session.commit()
            print(f"Fixed {len(fixed)} document lifecycle records (MISSING to PENDING_REVIEW)")

        # Fix: mark notifications as read for already-approved/rejected docs
        from sqlalchemy import text as sql_text
        try:
            db.session.execute(sql_text("""
                UPDATE notifications SET is_read = 1
                WHERE related_doc_id IN (
                    SELECT id FROM document_lifecycle
                    WHERE status IN ('APPROVED_AND_SIGNED', 'REVISION_REQUESTED')
                )
            """))
            db.session.commit()
            print("Marked notifications as read for processed documents")
        except Exception as e:
            db.session.rollback()
            print(f"Notification migration: {e}")

# Run setup unconditionally so Waitress initializes the DB tables
setup_app()

if __name__ == '__main__':
    # Start the application
    app.run(host='0.0.0.0', port=5055, debug=False)
